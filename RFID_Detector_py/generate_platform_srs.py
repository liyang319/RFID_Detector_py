"""生成 民用爆炸品企业服务管理平台 需求规格说明书"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def add_run(para, text, font_name='宋体', font_size=Pt(10.5), bold=False):
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    return run


def add_normal(doc, text, indent=False, bold=False):
    p = doc.add_paragraph()
    add_run(p, text, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    return p


def add_list(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(21)
    add_run(p, text)
    return p


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.font.bold = True


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.font.bold = True


def add_code(doc, text, indent=True):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if indent:
            p.paragraph_format.left_indent = Pt(42)
        add_run(p, line, font_name='Consolas', font_size=Pt(8.5))


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        add_run(cell.paragraphs[0], h, font_size=Pt(9), bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            add_run(cell.paragraphs[0], str(val), font_size=Pt(9))
    doc.add_paragraph()
    return table


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld)
        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        run2._r.append(instr)
        run3 = p.add_run()
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld2)


def build():
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.18)
        s.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tp, '民用爆炸品企业服务管理平台', font_name='黑体', font_size=Pt(18), bold=True)
    doc.add_paragraph()
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp, '需求说明书 v1.1', font_name='黑体', font_size=Pt(16), bold=True)
    doc.add_paragraph(); doc.add_paragraph()
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(vp, '版本：V1.1.0', font_size=Pt(12))
    doc.add_paragraph()
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(dp, f'日期：{datetime.date.today().isoformat()}', font_size=Pt(12))
    doc.add_page_break()

    # ===== 版本说明 =====
    add_normal(doc, '版本说明：', bold=True)
    add_normal(doc, '当前版本：V1.1.0，在V1.0.0基础上增加设备管理、标签打印管理、运输监控等功能。', indent=True)
    doc.add_paragraph()
    make_table(doc,
        ['序号', '变更章节', '变更说明', '变更日期'],
        [
            ['1', '全部', '初始版本', '2026-06-10'],
            ['2', '2,3,9,10,11', '增加设备管理、标签打印管理、运输监控；扩展数据来源为六种；更新数据字段定义', datetime.date.today().isoformat()],
        ])
    doc.add_page_break()

    # ===== 目录 =====
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tp, '目  录', font_name='黑体', font_size=Pt(14), bold=True)
    doc.add_paragraph()
    toc = [
        '1. 系统概述', '   1.1 软件标识', '   1.2 软件目标', '   1.3 硬件支撑', '   1.4 软件组成',
        '2. 用户权限管理', '   2.1 管理员权限', '   2.2 企业用户权限', '   2.3 监管用户权限',
        '3. 终端设备管理', '   3.1 设备类型', '   3.2 设备注册与认证', '   3.3 设备状态监控', '   3.4 设备OTA升级',
        '4. 标签打印内容管理', '   4.1 打印模板管理', '   4.2 打印信息下发', '   4.3 打印记录追溯',
        '5. 数据采集与接入', '   5.1 通道机数据上报（生产/入库/出库）', '   5.2 廊机数据上报', '   5.3 手持设备数据上报', '   5.4 销售端数据接入', '   5.5 爆破作业数据接入', '   5.6 靶试实验（科研）数据接入', '   5.7 运输数据接入', '   5.8 第三方平台数据对接',
        '6. 数据管理', '   6.1 产品全生命周期追踪', '   6.2 数据查询', '   6.3 数据统计', '   6.4 报表生成',
        '7. 生产管理', '   7.1 生产计划管理', '   7.2 生产数据接收', '   7.3 标签写入与校验记录',
        '8. 出入库管理', '   8.1 入库管理', '   8.2 出库管理', '   8.3 库存管理',
        '9. 销售管理', '   9.1 销售订单管理', '   9.2 销售数据上报',
        '10. 爆破作业管理', '   10.1 爆破作业登记', '   10.2 爆破数据上报',
        '11. 靶试实验管理',
        '12. 运输监控管理', '   12.1 运输计划管理', '   12.2 运输实时监控', '   12.3 运输轨迹回放', '   12.4 运输异常告警',
        '13. 系统对接', '   13.1 公安部管理平台对接', '   13.2 行业管理平台对接',
        '14. 告警与通知', '15. 系统设置', '16. 帮助',
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, item)
    doc.add_page_break()

    # ===== 1. 系统概述 =====
    add_h1(doc, '1. 系统概述')
    add_normal(doc, '软件名称：民用爆炸品企业服务管理平台')
    add_normal(doc, '软件目标：构建民用爆炸品从生产、入库、出库、运输、销售、爆破作业到靶试实验的全生命周期追溯与管控平台。平台管理各类终端设备（产线通道机、廊机、手持设备、数据大屏），接收设备上报的生产与出入库数据，下发标签打印信息和标签写入指令，对接销售端、爆破作业端及第三方监管平台，实现爆炸品"来源可查、去向可追、责任可究"的全程闭环管理。')
    add_normal(doc, '硬件支撑：云服务器、数据库服务器、产线通道机（RFID读写器+条码扫描器+IO模块）、廊机（库房门廊出入库识别设备）、手持终端（销售/爆破/运输数据采集）、数据大屏、网络设备等。')
    add_normal(doc, '软件组成：用户权限管理、终端设备管理、标签打印内容管理、数据采集与接入、数据管理、生产管理、出入库管理、运输监控管理、销售管理、爆破作业管理、靶试实验管理、系统对接、告警与通知、系统设置、帮助。')

    # ===== 2. 用户权限管理 =====
    add_h1(doc, '2. 用户权限管理')
    add_h2(doc, '2.1 管理员权限')
    add_list(doc, '全权限：包括用户管理、系统配置、数据删除、平台对接配置等所有功能；')
    add_list(doc, '可创建、编辑、删除企业用户账号和监管用户账号。')
    add_h2(doc, '2.2 企业用户权限')
    add_list(doc, '生产管理：查看和管理本企业生产计划、生产数据、标签写入记录；')
    add_list(doc, '出入库管理：查看和管理本企业入库、出库记录及库存数据；')
    add_list(doc, '销售管理：管理销售订单，上报销售数据；')
    add_list(doc, '数据查询：按条件查询本企业相关数据，查看统计报表；')
    add_list(doc, '爆破作业管理：查看本企业产品的爆破作业记录。')
    add_h2(doc, '2.3 监管用户权限')
    add_list(doc, '数据查看：查询所有企业上报的数据，按多维度检索；')
    add_list(doc, '统计报表：查看全域统计数据，生成监管报表；')
    add_list(doc, '告警接收：接收异常告警通知；')
    add_list(doc, '系统对接：管理第三方平台（公安部等）的对接配置。')

    # ===== 3. 终端设备管理 =====
    add_h1(doc, '3. 终端设备管理')
    add_normal(doc, '平台统一管理所有接入的终端设备，包括设备注册、认证、状态监控和OTA固件升级。')

    add_h2(doc, '3.1 设备类型')
    make_table(doc,
        ['设备类型', '应用场景', '主要功能', '通信协议'],
        [
            ['产线通道机', '安装在生产线末端', '标签写入、读取校验、条码扫描、数据上报', 'MQTT + TCP'],
            ['廊机', '安装在库房门廊处', '出入库识别、标签读取、条码扫描、数据上报', 'MQTT + TCP'],
            ['手持设备', '销售/爆破/运输现场', '标签读取、数据录入、定位上报', 'MQTT + HTTP REST'],
            ['数据大屏', '监控中心/调度室', '实时数据展示、告警展示、统计看板', 'HTTP REST / WebSocket'],
        ])

    add_h2(doc, '3.2 设备注册与认证')
    add_list(doc, '设备首次接入平台需进行注册，提交设备唯一标识（SN号/设备ID）、设备类型、所属企业/分厂、安装位置等信息；')
    add_list(doc, '平台审核通过后分配设备认证凭证（Token/证书），设备携带凭证进行后续通信；')
    add_list(doc, '支持设备信息编辑、禁用/启用、删除（注销）操作；')
    add_list(doc, '设备认证方式：基于Token的身份认证，每次MQTT连接和HTTP请求携带Token进行校验。')

    add_h2(doc, '3.3 设备状态监控')
    add_list(doc, '在线状态：实时展示各设备的在线/离线状态，离线超过阈值时间自动告警；')
    add_list(doc, '心跳检测：设备定时（默认30秒）向平台发送心跳消息，平台记录最近心跳时间；')
    add_list(doc, '设备信息查看：查看设备的基本信息、固件版本、最近上报数据时间、运行时长等；')
    add_list(doc, '设备列表筛选：按设备类型、所属企业、分厂、在线状态等条件筛选设备列表；')
    add_list(doc, '设备异常记录：记录设备异常事件（掉线、数据上报异常、硬件故障等），支持追溯查询。')

    add_h2(doc, '3.4 设备OTA升级')
    add_list(doc, '固件版本管理：平台管理各设备类型的固件版本，支持上传新版本固件包；')
    add_list(doc, '升级策略配置：支持立即升级、定时升级、分批次灰度升级；')
    add_list(doc, '升级目标选择：可按设备类型、所属企业、分厂、单个设备等多粒度选择升级目标；')
    add_list(doc, '升级进度监控：实时查看各设备的升级状态（等待中/下载中/安装中/成功/失败）；')
    add_list(doc, '升级失败回滚：升级失败时自动回滚到上一版本，并通知管理员；')
    add_list(doc, 'OTA升级Topic：firmware/upgrade/{device_id}（平台下发），firmware/status/{device_id}（设备上报状态）。')

    # ===== 4. 标签打印内容管理 =====
    add_h1(doc, '4. 标签打印内容管理')
    add_normal(doc, '平台统一管理不同产线的标签打印内容，支持按分厂、产线粒度定点下发打印信息。标签打印内容在贴标环节写入标签USER数据区，后续在通道机/廊机上被读取和校验。')

    add_h2(doc, '4.1 打印模板管理')
    add_list(doc, '模板创建：创建标签打印内容模板，定义各字段的内容和格式；')
    add_list(doc, '模板字段包括：产品类型、规格型号、生产厂家名称、生产厂家许可证号、生产日期、批号、包装规格、包装方式、数量、经纬度等；')
    add_list(doc, '模板支持变量替换：如{date}自动填充当前日期、{batch}自动填充批号等；')
    add_list(doc, '模板编辑与删除：支持模板的修改、复制、删除操作。')

    add_h2(doc, '4.2 打印信息下发')
    add_list(doc, '定点下发：选择目标分厂→产线，将指定的打印内容模板下发到对应产线的标签打印机或通道机；')
    add_list(doc, '下发内容格式：JSON格式，包含模板的所有字段值，通过MQTT下发到产线设备；')
    add_list(doc, '下发Topic：printer/content/{factory_id}/{line_id}（平台→设备）；')
    add_list(doc, '下发确认：设备收到打印信息后回复确认消息，平台记录下发状态；')
    add_list(doc, '批量下发：支持同时向多个分厂、多条产线下发相同或不同的打印内容；')
    add_list(doc, '下发历史：记录每次下发的目标、内容、时间、操作人，支持追溯查询。')
    add_normal(doc, '下发数据格式示例：', bold=True)
    print_fmt = '''{
    "type": "print_content",
    "factory_id": "F001",
    "line_id": "L03",
    "content": {
        "product_name": "乳化炸药",
        "package_spec": "Φ32mm×200g",
        "manufacturer": "XX化工有限公司",
        "license_number": "SC20240001",
        "production_date": "2026-06-11",
        "batch_number": "BATCH-20260611-001",
        "package_method": "箱装",
        "quantity": 24,
        "longitude": 116.3918,
        "latitude": 39.9798
    }
}'''
    add_code(doc, print_fmt)

    add_h2(doc, '4.3 打印记录追溯')
    add_list(doc, '每次打印下发自动记录：下发时间、目标分厂/产线、打印内容摘要、操作人员；')
    add_list(doc, '支持按时间范围、分厂、产线、产品类型等维度查询打印历史记录；')
    add_list(doc, '打印记录与后续生产数据通过EPC号关联，实现"打印→贴标→写入→校验"全链路追溯。')

    # ===== 5. 数据采集与接入 =====
    add_h1(doc, '5. 数据采集与接入')
    add_normal(doc, '平台数据来源分为六种：生产、入库、出库、销售、爆破作业、靶试实验（科研）。其中生产数据由产线通道机通过MQTT自动上报，入库和出库数据由廊机上报，销售和爆破作业数据通过手持设备或Web端录入上报，靶试实验数据通过专用接口上报，运输数据由手持设备或车载终端实时上报。')

    add_h2(doc, '5.1 产线通道机数据上报（生产）')
    add_normal(doc, '通道机（RFID标签识别系统）在完成标签写入、读取校验后，通过MQTT协议向平台上报数据。')
    add_normal(doc, 'MQTT上报Topic：rfid/data/{device_id}', bold=True)
    add_normal(doc, '上报数据格式：', bold=True)
    mqtt_fmt = '''{
    "cmd": "report_tags",
    "type": "inbound",
    "data": {
        "tags": [{
            "epc": "C090C000000A4B28",
            "tid": "",
            "user_data": "30 32 57 44 36 00 ...",
            "rssi": -42.5,
            "antenna_num": 1,
            "pc": "04 FF",
            "product_name": "乳化炸药",
            "manufacturer": "XX化工有限公司",
            "license_number": "SC20240001",
            "production_date": "2024-08-15",
            "batch_number": "BATCH-...",
            "package_spec": "标准规格",
            "package_method": "箱装",
            "quantity": 1,
            "longitude": 116.3918,
            "latitude": 39.9798,
            "timestamp": "2026-06-05 12:00:00",
            "write_verified": true
        }],
        "barcodes": ["6923456789012"],
        "validation": {
            "write_verified_count": 1,
            "write_total_count": 1,
            "errors": []
        },
        "write_success": true
    }
}'''
    add_code(doc, mqtt_fmt)

    add_normal(doc, '通道机也通过TCP发送实时消息（report_rfid / report_barcode / cargo_in / cargo_out / pass），平台可建立TCP连接接收。', bold=False)

    add_normal(doc, '生产端数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明', '来源'],
        [
            ['tid', 'string', '标签TID号', 'RFID读写器读取'],
            ['epc', 'string', '标签EPC号（唯一标识）', 'RFID读写器读取'],
            ['user_data', 'string', '用户数据区内容（十六进制）', '上位机下发→通道机写入'],
            ['longitude', 'float', '经度', '系统配置/通道机位置'],
            ['latitude', 'float', '纬度', '系统配置/通道机位置'],
            ['timestamp', 'datetime', '读取时间戳', '系统自动生成'],
            ['product_name', 'string', '产品类型（如乳化炸药）', '上位机下发'],
            ['package_spec', 'string', '规格型号', '上位机下发'],
            ['manufacturer', 'string', '生产厂家名称', '上位机下发'],
            ['license_number', 'string', '生产厂家许可证号', '上位机下发'],
            ['production_date', 'string', '生产日期', '上位机下发'],
            ['batch_number', 'string', '批号', '系统自动/上位机下发'],
            ['package_method', 'string', '包装方式', '上位机下发'],
            ['quantity', 'int', '数量', '上位机下发'],
            ['write_verified', 'bool', '标签写入校验是否通过', '通道机自动判定'],
        ])

    add_h2(doc, '5.2 廊机数据上报（入库/出库）')
    add_normal(doc, '廊机安装在库房门廊处，自动识别经过门廊的货物标签（入库或出库），并通过MQTT向平台上报数据。')
    add_list(doc, '廊机识别方向：通过内置光栅或红外传感器判断货物移动方向（入库/出库），与产线通道机状态机逻辑一致；')
    add_list(doc, '上报数据格式：与产线通道机相同，通过MQTT report_tags上报，data_type为"inbound"或"outbound"；')
    add_list(doc, '廊机也通过TCP发送实时消息（cargo_in / cargo_out / report_rfid / report_barcode），平台可建立TCP连接接收。')

    add_h2(doc, '5.3 手持设备数据接入')
    add_normal(doc, '手持设备用于销售现场、爆破作业现场、运输途中等场景的数据采集和上报。')
    add_list(doc, '功能包括：RFID标签读取、条码扫描、GPS定位、数据录入、拍照取证；')
    add_list(doc, '通信方式：通过4G/5G/WiFi网络，使用MQTT协议上报数据，使用HTTP REST进行配置同步；')
    add_list(doc, '上报数据类型：销售确认、爆破作业登记、运输状态更新、异常事件报告；')
    add_list(doc, '离线模式：网络不可用时本地缓存数据，恢复网络后自动补传。')

    add_h2(doc, '5.4 销售端数据接入')
    add_normal(doc, '销售端数据由销售企业在平台上录入或通过移动终端上报。')
    add_normal(doc, '销售端数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['epc', 'string', '产品EPC号'],
            ['longitude', 'float', '销售地点经度'],
            ['latitude', 'float', '销售地点纬度'],
            ['timestamp', 'datetime', '销售时间'],
            ['product_name', 'string', '产品类型'],
            ['package_spec', 'string', '规格型号'],
            ['sale_company', 'string', '销售公司名称'],
            ['quantity', 'int', '销售数量'],
            ['buyer_name', 'string', '购买方名称'],
            ['buyer_license', 'string', '购买方许可证号'],
        ])

    add_h2(doc, '5.5 爆破作业数据接入')
    add_normal(doc, '爆破作业数据由爆破公司在作业现场通过移动终端上报。')
    add_normal(doc, '爆破作业数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['epc', 'string', '产品EPC号'],
            ['longitude', 'float', '爆破作业地点经度'],
            ['latitude', 'float', '爆破作业地点纬度'],
            ['timestamp', 'datetime', '爆破作业时间'],
            ['product_name', 'string', '产品类型'],
            ['package_spec', 'string', '规格型号'],
            ['blast_company', 'string', '爆破公司名称'],
            ['blast_license', 'string', '爆破作业许可证号'],
            ['quantity', 'int', '爆破使用数量'],
            ['blast_method', 'string', '爆破方式'],
            ['blast_result', 'string', '爆破结果'],
            ['operator_name', 'string', '操作人员姓名'],
        ])

    add_h2(doc, '5.6 靶试实验（科研）数据接入')
    add_normal(doc, '科研/靶试实验数据由实验单位上报。')
    add_normal(doc, '科研数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['epc', 'string', '产品EPC号'],
            ['longitude', 'float', '实验地点经度'],
            ['latitude', 'float', '实验地点纬度'],
            ['timestamp', 'datetime', '实验时间'],
            ['product_name', 'string', '产品类型'],
            ['package_spec', 'string', '规格型号'],
            ['research_unit', 'string', '科研单位名称'],
            ['test_type', 'string', '实验类型（靶试/性能测试等）'],
            ['test_result', 'string', '实验结果'],
            ['report_file', 'string', '实验报告文件链接'],
        ])

    add_h2(doc, '5.7 运输数据接入')
    add_normal(doc, '运输过程中的数据由手持设备或车载终端实时上报，平台对运输全程进行监控。')
    add_normal(doc, '运输数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['transport_id', 'string', '运输任务编号'],
            ['epc_list', 'array', '运输产品EPC号列表'],
            ['vehicle_plate', 'string', '运输车辆车牌号'],
            ['driver_name', 'string', '驾驶员姓名'],
            ['driver_license', 'string', '驾驶员从业资格证号'],
            ['longitude', 'float', '当前位置经度（实时上报）'],
            ['latitude', 'float', '当前位置纬度（实时上报）'],
            ['speed', 'float', '当前速度（km/h）'],
            ['timestamp', 'datetime', '位置上报时间'],
            ['departure_place', 'string', '出发地'],
            ['destination', 'string', '目的地'],
            ['estimated_arrival', 'datetime', '预计到达时间'],
            ['transport_status', 'string', '运输状态：出发/在途/到达/异常'],
        ])
    add_list(doc, '运输数据上报Topic：transport/data/{transport_id}（MQTT）；')
    add_list(doc, '上报频率：正常行驶时每30秒上报一次位置，异常事件（停车超时、偏离路线等）立即上报。')

    add_h2(doc, '5.8 第三方平台数据对接')
    add_list(doc, '公安部管理平台对接：按《民用爆炸物品信息管理条例》要求，将产品生产、流通、使用数据同步上报至公安部管理平台。')
    add_list(doc, '行业管理平台对接：支持与行业协会或其他监管部门的平台进行数据交换。')
    add_list(doc, '对接方式：通过标准API接口（RESTful/WebService）或消息队列（MQTT/Kafka）进行数据同步。')
    add_list(doc, '对接数据格式：采用JSON格式，按对接平台要求进行字段映射和转换。')

    # ===== 4. 数据管理 =====
    add_h1(doc, '6. 数据管理')

    add_h2(doc, '6.1 产品全生命周期追踪')
    add_normal(doc, '以EPC号为唯一标识，串联产品从生产→入库→出库→销售→爆破作业（或靶试实验）的完整生命周期。')
    add_normal(doc, '全生命周期状态流转：', bold=True)
    add_list(doc, '生产下线 → 标签写入 → 入库 → 库存 → 出库 → 销售 → 爆破作业使用（或科研实验）')
    add_list(doc, '每个环节记录：EPC号、时间戳、经纬度、操作企业/单位、操作类型、操作人员。')

    add_h2(doc, '6.2 数据查询')
    add_normal(doc, '支持以下维度的数据查询：')
    add_list(doc, '按EPC号查询：输入单个EPC号，展示该产品的完整生命周期轨迹；')
    add_list(doc, '按时间范围查询：选择起止日期，查询该时段内所有环节的数据记录；')
    add_list(doc, '按产品类型查询：按产品名称（如乳化炸药、铵油炸药等）筛选；')
    add_list(doc, '按生产企业查询：按生产厂家名称或许可证号筛选；')
    add_list(doc, '按地理位置查询：按经纬度范围或行政区域筛选；')
    add_list(doc, '按批号查询：按生产批号查询该批次所有产品；')
    add_list(doc, '按操作类型查询：按生产/入库/出库/销售/爆破作业/靶试分类筛选；')
    add_list(doc, '组合查询：支持上述维度的任意组合，提供高级检索功能。')

    add_h2(doc, '6.3 数据统计')
    add_normal(doc, '支持以下维度的数据统计：')
    add_list(doc, '产量统计：按日/周/月/年统计各企业、各产品类型的生产数量；')
    add_list(doc, '库存统计：实时统计各企业、各仓库的当前库存量及分布；')
    add_list(doc, '销售统计：按时间段统计各销售公司的销售量和销售金额；')
    add_list(doc, '爆破使用统计：统计各爆破公司的使用量和爆破次数；')
    add_list(doc, '地域分布统计：以地图热力图展示产品的生产、存储、使用地理分布；')
    add_list(doc, '校验统计：统计标签写入校验的通过率、失败原因分析；')
    add_list(doc, '异常统计：统计异常告警的类型分布和处理情况。')

    add_h2(doc, '6.4 报表生成')
    add_list(doc, '支持按日/周/月/季/年生成各类统计报表；')
    add_list(doc, '报表类型：生产报表、库存报表、销售报表、爆破使用报表、全生命周期追溯报表、异常告警报表；')
    add_list(doc, '报表格式：PDF、Excel (.xlsx)、CSV；')
    add_list(doc, '支持报表模板自定义，用户可选择报表包含的字段和统计维度；')
    add_list(doc, '支持报表自动定时生成和邮件发送。')

    # ===== 5. 生产管理 =====
    add_h1(doc, '7. 生产管理')

    add_h2(doc, '7.1 生产计划管理')
    add_list(doc, '创建、编辑、删除生产计划，包括产品类型、规格型号、计划数量、计划日期、生产产线等；')
    add_list(doc, '生产计划审批流程：提交→审核→批准→执行；')
    add_list(doc, '生产计划执行状态跟踪：待执行/执行中/已完成/已取消。')

    add_h2(doc, '7.2 生产数据接收')
    add_normal(doc, '平台通过MQTT接收通道机上报的生产数据（type="inbound"或"outbound"，data_type参数区分入库出库，生产场景暂用inbound标识）。平台接收后自动完成：')
    add_list(doc, '数据解析与验证：校验数据格式完整性、必填字段是否缺失；')
    add_list(doc, '重复检测：基于EPC号+时间戳去重，防止重复上报；')
    add_list(doc, '数据入库存储：写入数据库持久化保存；')
    add_list(doc, '数据关联：将标签数据与生产计划、企业信息进行关联。')

    add_h2(doc, '7.3 标签写入与校验记录')
    add_list(doc, '记录每次标签写入操作的详细信息：写入数据、写入结果（成功/失败）、校验结果（匹配/不匹配）；')
    add_list(doc, '写入失败告警：当写入失败或校验不通过时，平台产生告警通知相关人员。')

    # ===== 6. 出入库管理 =====
    add_h1(doc, '8. 出入库管理')

    add_h2(doc, '8.1 入库管理')
    add_list(doc, '接收通道机上报的入库数据；')
    add_list(doc, '记录入库仓库、入库时间、操作人员；')
    add_list(doc, '入库数据与生产数据进行关联，标记产品状态为"已入库"；')
    add_list(doc, '支持批量入库操作和单条入库记录查询。')

    add_h2(doc, '8.2 出库管理')
    add_list(doc, '接收通道机上报的出库数据；')
    add_list(doc, '记录出库仓库、出库时间、操作人员、出库去向（销售/爆破作业/科研）；')
    add_list(doc, '出库数据与入库数据进行关联，标记产品状态为"已出库"；')
    add_list(doc, '支持出库审批流程：出库申请→审核→批准→出库执行。')

    add_h2(doc, '8.3 库存管理')
    add_list(doc, '实时库存查询：按仓库、产品类型、规格型号查看当前库存量；')
    add_list(doc, '库存预警：设置库存上限和下限阈值，超限时自动告警；')
    add_list(doc, '库存盘点：支持定期盘点功能，记录盘点差异；')
    add_list(doc, '库存台账：自动生成库存进出明细台账。')

    # ===== 7. 销售管理 =====
    add_h1(doc, '9. 销售管理')

    add_h2(doc, '9.1 销售订单管理')
    add_list(doc, '销售订单创建：选择产品类型、规格型号、数量、购买方信息；')
    add_list(doc, '销售订单审批：提交→审核→批准→发货；')
    add_list(doc, '购买方资质核验：核验购买方的民用爆炸物品购买许可证。')

    add_h2(doc, '9.2 销售数据上报')
    add_list(doc, '销售企业通过Web端或移动终端录入销售数据并上报平台；')
    add_list(doc, '销售数据与出库数据进行关联，标记产品状态为"已销售"；')
    add_list(doc, '销售数据上报Topic：sales/report/{company_id}（参考MQTT协议格式）。')

    # ===== 8. 爆破作业管理 =====
    add_h1(doc, '10. 爆破作业管理')

    add_h2(doc, '10.1 爆破作业登记')
    add_list(doc, '爆破公司在作业前通过平台登记爆破作业信息：作业地点、时间、使用产品类型和数量、操作人员等；')
    add_list(doc, '爆破作业审批：提交→审核→批准→执行；')
    add_list(doc, '核验爆破作业单位的许可证资质。')

    add_h2(doc, '10.2 爆破数据上报')
    add_list(doc, '爆破作业完成后，通过移动终端上报使用数据（EPC号、使用量、爆破结果等）；')
    add_list(doc, '爆破数据与出库/销售数据进行关联，标记产品状态为"已使用（爆破）"；')
    add_list(doc, '爆破数据不可修改或删除，确保数据真实性。')

    # ===== 9. 靶试实验管理 =====
    add_h1(doc, '11. 靶试实验管理')
    add_list(doc, '科研机构/企业在平台上登记靶试实验计划：实验类型、实验地点、时间、使用产品信息；')
    add_list(doc, '实验审批：提交→审核→批准→执行；')
    add_list(doc, '实验数据上报：实验完成后上报实验结果数据；')
    add_list(doc, '实验数据与出库数据进行关联，标记产品状态为"已使用（科研）"；')
    add_list(doc, '实验报告管理：支持上传和下载实验报告文件（PDF/Word/图片）。')

    # ===== 10. 系统对接 =====
    # ===== 12. 运输监控管理 =====
    add_h1(doc, '12. 运输监控管理')
    add_normal(doc, '平台对民用爆炸品的运输过程进行全程监控和管理，确保运输安全合规。')

    add_h2(doc, '12.1 运输计划管理')
    add_list(doc, '运输计划创建：指定运输任务编号、车辆信息、驾驶员信息、产品清单（EPC号列表）、出发地、目的地、预计路线、预计出发/到达时间；')
    add_list(doc, '运输计划审批：提交→审核→批准→执行；')
    add_list(doc, '运输资质核验：核验运输企业资质、车辆资质、驾驶员从业资格、押运员资质；')
    add_list(doc, '电子运单：生成电子运单，包含运输许可信息和产品清单，可打印或发送到手持设备。')

    add_h2(doc, '12.2 运输实时监控')
    add_list(doc, '实时位置追踪：在地图上实时显示运输车辆位置、速度、行驶方向；')
    add_list(doc, '产品状态监控：实时查看运输中产品的EPC列表和状态；')
    add_list(doc, '电子围栏：预设运输路线，车辆偏离路线超过设定距离（默认500米）时自动告警；')
    add_list(doc, '超时告警：车辆在非目的地停留超过设定时间（默认10分钟）时自动告警；')
    add_list(doc, '多车监控：支持同时监控多个运输任务，地图上同时显示多辆车的位置；')
    add_list(doc, '数据刷新频率：地图位置每30秒刷新一次，告警事件实时推送。')

    add_h2(doc, '12.3 运输轨迹回放')
    add_list(doc, '历史轨迹查询：选择运输任务和时间范围，在地图上回放运输轨迹；')
    add_list(doc, '轨迹回放控制：支持播放/暂停/快进/慢放，显示每个位置点的时间、速度信息；')
    add_list(doc, '停留点标注：自动标注运输途中的停留点及停留时长；')
    add_list(doc, '轨迹导出：支持导出运输轨迹数据为KML/GPX格式。')

    add_h2(doc, '12.4 运输异常告警')
    make_table(doc,
        ['告警类型', '触发条件', '告警级别', '处置方式'],
        [
            ['路线偏离', '偏离预设路线超过阈值', '高', '平台弹窗 + 短信通知驾驶员和管理员'],
            ['超时停留', '非目的地停留超过阈值', '高', '平台弹窗 + 短信通知驾驶员'],
            ['超速行驶', '超过道路限速或设定的最高速度', '中', '平台弹窗'],
            ['设备离线', '车载终端/手持设备心跳超时', '高', '平台弹窗 + 短信通知管理员'],
            ['到达确认', '车辆到达目的地并确认产品交接', '信息', '更新运输状态为"已完成"'],
            ['异常开箱', '运输途中异常打开产品包装（如有传感器）', '高', '平台弹窗 + 短信/电话通知管理员'],
        ])

    # ===== 13. 系统对接 =====
    add_h1(doc, '13. 系统对接')

    add_h2(doc, '13.1 公安部管理平台对接')
    add_normal(doc, '按照GA/T 1225-2015《民用爆炸物品信息管理系统数据交换标准》等相关行业标准，实现与公安部民爆物品管理平台的数据对接。')
    add_list(doc, '数据同步内容：生产数据、出入库数据、销售数据、爆破作业数据；')
    add_list(doc, '数据同步方式：通过公安部指定的数据交换接口（API/中间库）定时或实时同步；')
    add_list(doc, '数据同步频率：支持实时同步和定时批量同步（可配置间隔）；')
    add_list(doc, '同步状态监控：记录每次同步的时间、数据量、成功/失败状态，失败时自动重试并告警。')

    add_h2(doc, '13.2 行业管理平台对接')
    add_list(doc, '支持与省级/市级民爆行业监管平台的数据对接；')
    add_list(doc, '支持与企业ERP/WMS系统的数据集成；')
    add_list(doc, '提供标准RESTful API接口供第三方系统调用查询。')

    # ===== 11. 告警与通知 =====
    add_h1(doc, '14. 告警与通知')
    add_normal(doc, '平台对以下异常情况进行告警：')
    make_table(doc,
        ['告警类型', '触发条件', '告警级别', '通知方式'],
        [
            ['标签写入失败', '通道机上报write_success=false', '高', '平台弹窗 + 短信/邮件通知'],
            ['标签校验失败', 'write_verified_count < write_total_count', '高', '平台弹窗 + 短信/邮件通知'],
            ['数量不一致', '标签数量 ≠ 条码数量', '中', '平台弹窗'],
            ['库存超限', '库存低于下限或高于上限', '中', '平台弹窗 + 邮件通知'],
            ['超时未上报', '通道机超过设定时间无数据上报', '中', '平台弹窗 + 短信通知'],
            ['数据异常', '关键字段缺失或格式错误', '低', '平台弹窗'],
            ['数据上报时间异常', '数据上报时间与实际时间偏差过大', '低', '平台弹窗'],
            ['第三方同步失败', '向公安部/行业平台同步数据失败', '高', '平台弹窗 + 短信通知'],
            ['未审批操作', '无审批的出入库/销售/爆破操作', '高', '平台弹窗 + 短信/邮件通知'],
        ])
    add_normal(doc, '告警记录管理：所有告警记录保存至数据库，支持按时间、类型、级别查询；告警处理流程：产生→确认→处理→关闭，记录处理人和处理时间。')

    # ===== 12. 系统设置 =====
    add_h1(doc, '15. 系统设置')
    add_list(doc, '用户管理：创建、编辑、删除用户账号，分配角色权限；')
    add_list(doc, '企业管理：管理生产企业、销售公司、爆破公司的基本信息（名称、许可证号、地址等）；')
    add_list(doc, '仓库管理：管理仓库信息（名称、位置、容量等）；')
    add_list(doc, '产品类型管理：管理产品分类和规格型号字典；')
    add_list(doc, '告警规则配置：配置各类告警的触发条件、级别和通知方式；')
    add_list(doc, '数据对接配置：配置与公安部平台、行业平台的对接参数（接口地址、认证信息、同步频率等）；')
    add_list(doc, '界面设置：界面风格、字体、主题切换；')
    add_list(doc, '系统日志：记录用户操作日志和系统运行日志，支持日志查询和导出。')

    # ===== 13. 帮助 =====
    add_h1(doc, '16. 帮助')
    add_list(doc, '版本说明：当前版本V1.1.0，后续版本更新记录；')
    add_list(doc, '使用说明：平台各模块的操作手册，含图文教程；')
    add_list(doc, 'API文档：平台对外接口的技术文档，含请求/响应示例；')
    add_list(doc, '常见问题（FAQ）：常见问题及解答；')
    add_list(doc, '在线支持：联系方式、工单提交。')

    add_page_numbers(doc)

    path = '民用爆炸品企业服务管理平台需求说明书 v1.1.docx'
    doc.save(path)
    print(f'文档已生成: {path}')


if __name__ == '__main__':
    build()
