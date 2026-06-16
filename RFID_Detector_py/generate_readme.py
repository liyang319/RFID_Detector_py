"""生成 RFID标签识别系统 软件功能说明文档 (Word格式)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    return h


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # ===== 封面 / 标题 =====
    title = doc.add_heading('RFID标签识别系统', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('软件功能说明文档')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'版本: v1.0.0\n').font.size = Pt(10)
    meta.add_run(f'生成日期: {datetime.date.today().isoformat()}\n').font.size = Pt(10)
    meta.add_run(f'设备ID: RFID-DETECTOR-001').font.size = Pt(10)

    doc.add_page_break()

    # ===== 目录占位 =====
    doc.add_heading('目  录', level=1)
    toc_items = [
        '1. 系统概述',
        '2. 软件架构',
        '3. 系统架构图',
        '4. 模块详细说明',
        '   4.1 main.py — 主控模块',
        '   4.2 RFIDReader_SFM2200.py — RFID读写器驱动',
        '   4.3 serial_comm.py — IO串口通信模块',
        '   4.4 barcode_scanner.py — 条码扫描器模块',
        '   4.5 mqtt_client.py — MQTT通信模块',
        '   4.6 TcpSocketServer.py — TCP服务端模块',
        '   4.7 rfid_tag.py — RFID标签数据模型',
        '   4.8 command.py — 指令定义模块',
        '5. 业务流程',
        '   5.1 入库流程',
        '   5.2 出库流程',
        '6. 数据协议',
        '   6.1 RFID标签上报数据格式',
        '   6.2 TCP下发写入数据格式',
        '   6.3 TCP通行完成消息格式',
        '   6.4 MQTT上报数据格式',
        '7. 硬件接口',
        '8. UI界面说明',
        '9. 部署说明',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    doc.add_page_break()

    # ===== 1. 系统概述 =====
    add_heading_styled(doc, '1. 系统概述', 1)
    doc.add_paragraph(
        'RFID标签识别系统是一套用于仓库出入库管理的自动化识别软件。'
        '系统通过串口连接RFID读写器(SFM2200)、IO模块(光栅/指示灯)和条码扫描器，'
        '实现货物标签的自动写入、读取、校验与数据上报。'
    )
    doc.add_paragraph('主要功能包括：')
    features = [
        'RFID标签自动写入与读取验证',
        '光栅状态机实现入库/出库方向自动识别',
        '条码扫描与RFID标签数量比对',
        '通过MQTT向主机上报标签数据和校验结果',
        '通过TCP Socket Server接收主机下发的写入数据',
        '通过TCP向主机反馈通行完成消息',
        '工业风格Tkinter图形界面，支持实时状态监控',
        '指示灯控制（红/黄/绿）实现状态可视化',
        '数据导出到Excel文件',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # ===== 2. 软件架构 =====
    add_heading_styled(doc, '2. 软件架构', 1)
    doc.add_paragraph(
        '系统采用多线程架构，主线程负责Tkinter GUI，各功能模块运行在独立的后台线程中，'
        '通过队列和回调机制进行线程间通信。UI更新统一通过 root.after(0, callback) 实现线程安全。'
    )

    add_heading_styled(doc, '2.1 关键技术栈', 2)
    make_table(doc,
               ['层级', '技术', '说明'],
               [
                   ['UI框架', 'Tkinter', 'Python标准GUI库，工业风格配色'],
                   ['串口通信', 'pyserial', 'RFID读写器、IO模块、条码扫描器'],
                   ['网络通信', 'socket (TCP)', 'TCP Server接收主机下发数据'],
                   ['MQTT', 'paho-mqtt', '标签数据上报到MQTT Broker'],
                   ['数据导出', 'openpyxl', '标签数据导出为Excel (.xlsx)'],
                   ['线程模型', 'threading', '多线程并发处理'],
               ])

    add_heading_styled(doc, '2.2 线程模型', 2)
    make_table(doc,
               ['线程', '职责', '启动方式'],
               [
                   ['主线程 (Tkinter)', 'GUI事件循环，UI渲染', '程序入口'],
                   ['IO串口读取线程', '轮询光栅状态，驱动状态机', 'auto_connect() → 后台daemon线程'],
                   ['RFID串口接收线程', '接收SFM2200标签上报数据', 'RFIDReader_SFM2200.start_receive_loop()'],
                   ['条码扫描器接收线程', '接收条码扫描器数据', 'BarCodeScanner.start_receive_loop()'],
                   ['MQTT网络循环', 'paho-mqtt内置loop线程', 'mqtt_client.connect()'],
                   ['TCP Server accept线程', '监听并接受TCP客户端连接', 'TcpSocketServer.start()'],
                   ['TCP 客户端处理线程', '每个TCP客户端独立收发线程', 'accept时动态创建'],
                   ['自动连接线程', '延迟初始化各模块', 'auto_connect() → 多个daemon线程'],
               ])

    # ===== 3. 系统架构图 =====
    add_heading_styled(doc, '3. 系统架构图', 1)

    doc.add_paragraph('以下为系统的整体架构示意（文本图）：')

    arch_text = """
┌─────────────────────────────────────────────────────────────────┐
│                        RFID标签识别系统                          │
│                     (Tkinter GUI 主线程)                         │
├─────────────────────────────────────────────────────────────────┤
│  ┌──────────┐ ┌───────────┐ ┌──────────┐ ┌──────────────────┐  │
│  │ 数据看板  │ │ 取标内容   │ │ 系统日志  │ │ 控制面板(运行/停止)│  │
│  └──────────┘ └───────────┘ └──────────┘ └──────────────────┘  │
├─────────────────────────────────────────────────────────────────┤
│                        业务逻辑层                                │
│  ┌──────────────────────┐  ┌─────────────────────────────────┐  │
│  │   光栅状态机           │  │   标签写入/读取/校验流水线       │  │
│  │   IDLE → START →      │  │   TCP收数据 → 写标签 → 读标签   │  │
│  │   MIDDLE → END → IDLE │  │   → 校验 → MQTT上报 → TCP回传   │  │
│  └──────────────────────┘  └─────────────────────────────────┘  │
├─────────────────────────────────────────────────────────────────┤
│                        数据模型层                                │
│  ┌──────────────────────────────────────────────────────────┐   │
│  │  RFIDTag (rfid_tag.py) — EPC/TID/USER/RSSI/产品信息       │   │
│  └──────────────────────────────────────────────────────────┘   │
├─────────────────────────────────────────────────────────────────┤
│                      硬件通信驱动层                               │
│  ┌──────────┐ ┌────────────┐ ┌───────────┐ ┌──────────────┐   │
│  │RFIDReader │ │ SerialComm │ │BarCodeScan │ │ TcpSocket     │   │
│  │_SFM2200  │ │ (IO模块)    │ │ (条码扫描) │ │ Server (TCP)  │   │
│  │(RFID读写) │ │ 光栅/灯控制 │ │ 条码读取   │ │ 上位机通信    │   │
│  └──────────┘ └────────────┘ └───────────┘ └──────────────┘   │
├─────────────────────────────────────────────────────────────────┤
│                      外部通信层                                  │
│  ┌──────────┐  ┌───────────────────────────────────────────┐   │
│  │MQTT Broker│  │           上位机系统                       │   │
│  │(数据上报) │  │  (TCP下发写入数据 + 接收pass消息)          │   │
│  └──────────┘  └───────────────────────────────────────────┘   │
├─────────────────────────────────────────────────────────────────┤
│                        硬件设备层                                │
│  ┌────────┐ ┌────────┐ ┌────────┐ ┌──────┐ ┌──────┐          │
│  │RFID读写│ │光栅×2  │ │指示灯  │ │条码枪 │ │蜂鸣器│          │
│  │SFM2200│ │(入口/出)│ │(红黄绿)│ │      │ │      │          │
│  └────────┘ └────────┘ └────────┘ └──────┘ └──────┘          │
└─────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, arch_text)

    doc.add_paragraph()
    doc.add_paragraph('数据流示意图：')

    dataflow_text = """
  [上位机]                    [RFID标签识别系统]                   [设备层]
     │                              │                               │
     │──TCP JSON(写入数据)──→   TCP Server                          │
     │                              │                               │
     │                         状态机(START)                         │
     │                              │──→ write_tag_with_userdata() ─→ [RFID读写器]
     │                              │                                   │ 写入标签
     │                              │←── startloop() 读取验证 ←─────────┘
     │                              │
     │                              │  (货物通过光栅)
     │                              │←── 光栅状态变化 ─────────────── [IO模块]
     │                              │
     │                              │  (END状态: 校验+上报)
     │                              │──→ MQTT report_tags ──────────→ [MQTT Broker]
     │                              │──→ TCP pass消息 ──────────────→ [上位机]
     │                              │
     │                              │←── 条码数据 ────────────────── [条码扫描器]
"""
    add_code_block(doc, dataflow_text)

    # ===== 4. 模块详细说明 =====
    add_heading_styled(doc, '4. 模块详细说明', 1)

    # --- 4.1 main.py ---
    add_heading_styled(doc, '4.1 main.py — 主控模块', 2)
    doc.add_paragraph(
        '主控模块，包含RFIDProductionSystem类，是整个系统的核心。负责GUI创建、'
        '各模块初始化协调、状态机流转、标签写入/读取流水线控制、MQTT上报和TCP消息处理。'
    )
    add_heading_styled(doc, '核心类', 3)
    make_table(doc,
               ['类名', '说明'],
               [
                   ['RFIDProductionSystem', '系统主类，继承object，管理所有子模块和UI'],
               ])

    add_heading_styled(doc, '关键属性', 3)
    make_table(doc,
               ['属性', '类型', '说明'],
               [
                   ['serial_comm', 'SerialComm', 'IO模块串口通信（光栅状态读取、指示灯控制）'],
                   ['rfid_reader_serial', 'RFIDReader_SFM2200', '串口RFID读写器（标签读写）'],
                   ['bar_scanner', 'BarCodeScanner', '条码扫描器'],
                   ['mqtt_client', 'MqttClient', 'MQTT客户端（数据上报）'],
                   ['tcp_server', 'TcpSocketServer', 'TCP服务端（接收上位机数据、回传pass消息）'],
                   ['tag_history', 'list[RFIDTag]', '当前批次标签历史记录'],
                   ['pending_write_data', 'bytes/None', 'TCP下发的待写入数据'],
                   ['actual_write_data', 'bytes/None', '实际写入的数据（用于校验）'],
                   ['write_done', 'bool', '写入是否完成标志'],
                   ['write_in_progress', 'bool', '写入是否进行中'],
                   ['direction', 'int', '方向标志：0=无, 1=入库, 2=出库'],
                   ['current_status', 'int', '当前光栅状态（bit位编码）'],
               ])

    add_heading_styled(doc, '关键方法', 3)
    make_table(doc,
               ['方法', '说明'],
               [
                   ['start_serial_reading_loop()', 'IO串口读取循环（状态机主循环），后台线程运行'],
                   ['_execute_fixed_write()', '执行标签写入（优先TCP数据，fallback默认数据），含重试'],
                   ['on_tcp_message(data, addr)', 'TCP消息回调，解析JSON写入数据或控制命令'],
                   ['_send_tcp_pass_message()', '发送 {"type":"pass","number":N} 给TCP客户端'],
                   ['report_rfid_tags_via_mqtt()', '通过MQTT上报标签数据和校验结果'],
                   ['_add_serial_tag_to_history(tag)', '标签添加到历史记录（EPC去重，写入后覆盖USER_DATA）'],
                   ['_display_tags_in_fetch()', '将当前标签显示在取标内容窗口'],
                   ['toggle_production()', '手动运行/停止切换，控制黄/绿灯和RFID读取'],
                   ['emergency_stop()', '手动停止（紧急制动），停止RFID读取，显示标签，上报数据'],
                   ['start_rfid_loop_query(b_on)', '出入库开始/结束时控制指示灯和RFID读取'],
                   ['update_ui_with_reported_tags()', '更新取标内容UI（含校验✓/✗标记）'],
               ])

    add_heading_styled(doc, '状态机（光栅状态流转）', 3)
    state_text = """
状态定义:
  STATE_IDLE = 0            (空闲)
  STATE_INBOUND_START = 1   (入库开始 — 光栅1遮挡)
  STATE_INBOUND_MIDDLE = 2  (入库中间 — 光栅1+2同时遮挡)
  STATE_INBOUND_END = 3     (入库结束 — 光栅2遮挡)
  STATE_OUTBOUND_START = 4  (出库开始 — 光栅2遮挡)
  STATE_OUTBOUND_MIDDLE = 5 (出库中间 — 光栅1+2同时遮挡)
  STATE_OUTBOUND_END = 6    (出库结束 — 光栅1遮挡)

光栅状态编码:
  0x01 = 光栅1遮挡 (入口端)
  0x02 = 光栅2遮挡 (出口端)
  0x03 = 光栅1+2同时遮挡

入库Path1: IDLE →(0x01)→ START →(0x03)→ MIDDLE →(0x02)→ END →(0x00)→ IDLE
入库Path2: IDLE →(0x01)→ START →(0x00)→ END →(0x00)→ IDLE  (跳过MIDDLE)
出库Path1: IDLE →(0x02)→ START →(0x03)→ MIDDLE →(0x01)→ END →(0x00)→ IDLE
出库Path2: IDLE →(0x02)→ START →(0x00)→ END →(0x00)→ IDLE  (跳过MIDDLE)
"""
    add_code_block(doc, state_text)

    add_heading_styled(doc, '标签写入流水线', 3)
    pipeline_text = """
START状态触发:
  1. 清空tag_history, 清空条码缓存
  2. 设置 write_done=False, write_in_progress=False
  3. 调用 _execute_fixed_write()
     ├── write_data = pending_write_data (TCP下发) 或 FIXED_USER_DATA (默认)
     ├── 记录 actual_write_data 用于后续校验
     ├── RFIDReader_SFM2200.write_tag_with_userdata(write_data)
     │   └── 失败时自动重试1次
     └── 成功后调用 startloop() 启动标签读取验证

报告时(END状态/MIDDLE快速通过/手动停止):
  1. 比较 read_user_data vs actual_write_data → 校验结果(✓/✗)
  2. MQTT上报标签数据+条码+校验结果
  3. TCP发送 {"type":"pass","number":N}
  4. 清除tag_history, write_done=False, actual_write_data=None
     (pending_write_data保留, 等待新TCP数据覆盖)
"""
    add_code_block(doc, pipeline_text)

    # --- 4.2 RFIDReader_SFM2200.py ---
    add_heading_styled(doc, '4.2 RFIDReader_SFM2200.py — RFID读写器驱动', 2)
    doc.add_paragraph(
        '封装SFM2200 RFID读写器的串口通信协议，支持标签读取和写入。'
        '包含CRC校验、指令发送/响应接收、标签数据解析等功能。'
    )
    make_table(doc,
               ['方法', '说明'],
               [
                   ['open() / close()', '打开/关闭串口连接'],
                   ['start_firmware()', '启动固件'],
                   ['startloop()', '发送标签盘存（读取）指令'],
                   ['stoploop()', '发送停止盘存指令'],
                   ['write_tag_with_userdata(userdata)', '向标签写入20字节USER数据，返回bool'],
                   ['set_callback(func)', '设置标签上报数据回调'],
                   ['set_write_callback(func)', '设置写入结果回调'],
                   ['calc_crc(data)', '计算CRC16校验值'],
                   ['_receive_loop()', '后台接收线程，自动区分标签上报和指令响应'],
                   ['_is_tag_report(data)', '判断数据是否为标签主动上报（FF xx AA开头）'],
               ])

    add_heading_styled(doc, '标签上报数据格式', 3)
    tag_fmt = """
  字节0-6:   固定头 FF 2B AA 00 00 00 96 (7字节)
  字节7:     RSSI (有符号)
  字节8:     天线号
  字节9-12:  时间戳
  字节13-14: Tag Data Length
  字节15-34: User Data (20字节)
  字节35:    EPC Length
  字节36-37: PC
  字节38+:   EPC数据 (长度 = EPC Length - 4)
"""
    add_code_block(doc, tag_fmt)

    # --- 4.3 serial_comm.py ---
    add_heading_styled(doc, '4.3 serial_comm.py — IO串口通信模块', 2)
    doc.add_paragraph(
        '封装IO模块的串口通信，采用类Modbus协议（FE + 命令 + 数据 + CRC16）。'
        '用于读取光栅状态和控制指示灯。'
    )
    make_table(doc,
               ['方法', '说明'],
               [
                   ['read_register(cmd, timeout)', '读取寄存器数据（如光栅状态 0x02寄存器）'],
                   ['write_register(port, b_on, timeout)', '写寄存器（控制灯：0x00红, 0x02黄, 0x04绿, 0x06蜂鸣器），b_on=True=0xFF亮, b_on=False=0x00灭'],
                   ['crc16(data, length)', 'CRC16/MODBUS校验计算'],
                   ['send(data) / receive()', '底层串口收发'],
               ])

    # --- 4.4 barcode_scanner.py ---
    add_heading_styled(doc, '4.4 barcode_scanner.py — 条码扫描器模块', 2)
    doc.add_paragraph(
        '封装条码扫描器的串口通信，后台线程持续接收条码数据，'
        '支持条码队列管理、去重、CSV日志记录。'
    )
    make_table(doc,
               ['方法', '说明'],
               [
                   ['start_receive_loop()', '启动后台接收线程'],
                   ['get_all_barcodes()', '获取并清空当前累积的所有条码'],
                   ['get_barcode()', '获取单个条码（FIFO）'],
                   ['add_barcode(barcode)', '添加条码到队列（自动去重）'],
                   ['log_barcode(barcode)', '记录条码到CSV日志文件'],
                   ['set_callback(func)', '设置条码接收回调'],
               ])

    # --- 4.5 mqtt_client.py ---
    add_heading_styled(doc, '4.5 mqtt_client.py — MQTT通信模块', 2)
    doc.add_paragraph(
        '封装paho-mqtt客户端，提供连接/发布/订阅/消息队列功能。'
        '用于向主机系统上报RFID标签数据。'
    )
    make_table(doc,
               ['属性/方法', '说明'],
               [
                   ['data_topic', '"rfid/command/{device_id}"'],
                   ['response_topic', '"rfid/response/{device_id}"'],
                   ['command_topic', '"rfid/data/{device_id}" — 数据上报主题'],
                   ['connect()', '连接Broker并启动loop'],
                   ['publish(topic, msg)', '发布MQTT消息'],
                   ['subscribe(topic)', '订阅主题'],
                   ['get_message()', '从消息队列获取消息（非阻塞）'],
                   ['mqtt_report_rfid_tags()', '发送report_tags指令'],
               ])

    # --- 4.6 TcpSocketServer.py ---
    add_heading_styled(doc, '4.6 TcpSocketServer.py — TCP服务端模块', 2)
    doc.add_paragraph(
        '多线程TCP Socket Server，监听端口3000，支持多个客户端并发连接。'
        '用于接收上位机下发的写入数据（JSON格式），并向客户端广播消息。'
    )
    make_table(doc,
               ['方法', '说明'],
               [
                   ['start() / stop()', '启动/停止服务器'],
                   ['register_callback(cb)', '注册消息接收回调 cb(data:bytes, addr:tuple)'],
                   ['send_to_all(msg)', '向所有已连接客户端广播消息（str或bytes）'],
                   ['_handle_client(socket, addr)', '独立线程处理单客户端收发'],
               ])

    # --- 4.7 rfid_tag.py ---
    add_heading_styled(doc, '4.7 rfid_tag.py — RFID标签数据模型', 2)
    doc.add_paragraph(
        'RFIDTag类，定义标签的完整数据结构。支持从原始字节解析标签数据、'
        '与字典互相转换、格式化摘要输出。'
    )
    make_table(doc,
               ['字段', '类型', '说明'],
               [
                   ['epc', 'str', 'EPC数据（十六进制字符串）'],
                   ['tid', 'str', 'TID数据'],
                   ['user_data', 'str', 'USER数据（空格分隔十六进制）'],
                   ['rssi', 'float', '信号强度（dBm）'],
                   ['antenna_num', 'int', '天线号'],
                   ['pc', 'str', 'PC数据'],
                   ['product_name', 'str', '产品名称'],
                   ['manufacturer', 'str', '生产企业'],
                   ['license_number', 'str', '生产许可证编号'],
                   ['production_date', 'str', '生产日期'],
                   ['batch_number', 'str', '批号'],
                   ['package_spec', 'str', '包装规格'],
                   ['package_method', 'str', '包装方式'],
                   ['quantity', 'int', '数量'],
                   ['longitude / latitude', 'float', '经纬度'],
                   ['timestamp', 'str', '读取时间戳'],
                   ['success', 'bool', '解析是否成功'],
               ])

    # --- 4.8 command.py ---
    add_heading_styled(doc, '4.8 command.py — 指令定义模块', 2)
    doc.add_paragraph(
        '定义各类十六进制指令常量（当前仅保留RFID Loop Start/Stop指令，'
        '用于兼容旧版CNNT读写器，该读写器已注释停用）。'
    )

    # ===== 5. 业务流程 =====
    add_heading_styled(doc, '5. 业务流程', 1)

    add_heading_styled(doc, '5.1 入库流程', 2)
    inbound_text = """
1. 上位机通过TCP发送写入数据: {"type":"rfid","data":[...]}
2. 货物从入口方向进入通道，光栅1被遮挡 → 状态机进入 STATE_INBOUND_START
3. START状态:
   - 黄灯亮，绿灯灭
   - 清空tag_history和条码缓存
   - 调用 _execute_fixed_write() 写入标签（优先使用TCP下发的pending_write_data）
   - RFID读写器开始盘存（读取标签验证）
4. 货物继续前进:
   Path1: 光栅1+2同时遮挡 → STATE_INBOUND_MIDDLE (如write未触发则补偿写入)
          → 光栅1松开,仅光栅2遮挡 → STATE_INBOUND_END
   Path2: 直接无遮挡 → STATE_INBOUND_END (跳过MIDDLE)
5. END状态 (光栅2松开,恢复无遮挡):
   - 绿灯亮，黄灯灭
   - 停止RFID盘存
   - 获取所有条码
   - 校验: 比较读回的user_data与actual_write_data (✓/✗)
   - MQTT上报: report_tags (含标签数据、条码、校验结果)
   - TCP回传: {"type":"pass","number":N}
   - 清空tag_history，重置write_done
   - 状态机回到 IDLE
"""
    add_code_block(doc, inbound_text)

    add_heading_styled(doc, '5.2 出库流程', 2)
    outbound_text = """
1. 上位机通过TCP发送写入数据
2. 货物从出口方向进入通道，光栅2被遮挡 → 状态机进入 STATE_OUTBOUND_START
3. START状态: 同入库 (黄灯亮, 写标签, 开始盘存)
4. 货物继续前进:
   Path1: 光栅1+2同时遮挡 → STATE_OUTBOUND_MIDDLE
          → 光栅1遮挡 → STATE_OUTBOUND_END
   Path2: 直接无遮挡 → STATE_OUTBOUND_END
5. END状态: 同入库 (绿灯亮, 校验, MQTT上报, TCP回传, 回到IDLE)
"""
    add_code_block(doc, outbound_text)

    # ===== 6. 数据协议 =====
    add_heading_styled(doc, '6. 数据协议', 1)

    add_heading_styled(doc, '6.1 RFID标签上报数据格式', 2)
    doc.add_paragraph('SFM2200读写器上报的标签数据，以 FF 2B AA 00 00 00 96 为包头。参见4.2节。')

    add_heading_styled(doc, '6.2 TCP下发写入数据格式', 2)
    doc.add_paragraph('上位机 → TCP Server (端口3000)，JSON格式：')
    tcp_data_fmt = '''{
    "type": "rfid",
    "data": [48, 50, 87, 68, 54, 0, 11, 0, 24, 88, 0, 12, 26, 1, 1, 0, 100, 0, 10]
}
说明:
  - "type": "rfid" 标识为RFID写入数据
  - "data": 整数数组, 每个值0-255, 系统自动转为bytes并补齐/截断到20字节
  - 不足20字节时尾部补0x00, 超过20字节时截取前20字节'''
    add_code_block(doc, tcp_data_fmt)

    add_heading_styled(doc, '6.3 TCP通行完成消息格式', 2)
    doc.add_paragraph('系统 → 上位机（TCP广播），JSON格式：')
    pass_fmt = '''{
    "type": "pass",
    "number": 5
}
说明:
  - "type": "pass" 标识为通行完成
  - "number": 本次识别到的RFID标签数量'''
    add_code_block(doc, pass_fmt)

    add_heading_styled(doc, '6.4 MQTT上报数据格式', 2)
    doc.add_paragraph('系统 → MQTT Broker → 主机，发布到 rfid/data/{device_id}：')
    mqtt_fmt = '''{
    "cmd": "report_tags",
    "type": "inbound",       // "inbound"(入库) 或 "outbound"(出库)
    "data": {
        "tags": [
            {
                "epc": "C090C000000A4B28",
                "tid": "",
                "user_data": "30 32 57 44 36 00 0B 00 ...",
                "rssi": -42.5,
                "antenna_num": 1,
                "product_name": "乳化炸药",
                ...
                "write_verified": true    // true=写入校验通过, false=不通过
            }
        ],
        "barcodes": ["6923456789012", ...],
        "validation": {
            "write_verified_count": 1,
            "write_total_count": 1,
            "errors": []
        },
        "write_success": true
    }
}'''
    add_code_block(doc, mqtt_fmt)

    # ===== 7. 硬件接口 =====
    add_heading_styled(doc, '7. 硬件接口', 1)
    make_table(doc,
               ['设备', '接口类型', '端口(默认)', '协议'],
               [
                   ['RFID读写器 (SFM2200)', '串口 (USB转TTL)', '/dev/ttysWK3', 'SFM2200私有协议+CRC'],
                   ['IO模块 (光栅+指示灯+蜂鸣器)', '串口 (RS485/RS232)', '/dev/ttyS0', '类Modbus RTU (FE+CMD+CRC16)'],
                   ['条码扫描器', '串口 (USB转TTL)', '/dev/ttyS1', 'ASCII字符串'],
                   ['上位机TCP通信', '以太网', '端口3000', 'TCP Socket + JSON'],
                   ['MQTT Broker', '以太网', '192.168.3.83:1883', 'MQTT 3.1.1'],
               ])

    add_heading_styled(doc, 'IO模块寄存器映射', 2)
    make_table(doc,
               ['地址', '功能', '说明'],
               [
                   ['0x00', '红灯控制', '0xFF=亮, 0x00=灭'],
                   ['0x02', '黄灯控制', '0xFF=亮, 0x00=灭'],
                   ['0x04', '绿灯控制', '0xFF=亮, 0x00=灭'],
                   ['0x06', '蜂鸣器控制', '0xFF=响, 0x00=停'],
                   ['0x02 (读)', '光栅状态读取', 'bit0=光栅1, bit1=光栅2'],
               ])

    # ===== 8. UI界面说明 =====
    add_heading_styled(doc, '8. UI界面说明', 1)
    doc.add_paragraph('系统采用Tkinter工业风格配色（深蓝主背景 + 浅灰面板），支持滚动。界面分为以下区域：')

    make_table(doc,
               ['区域', '内容', '说明'],
               [
                   ['标题栏', '"RFID标签识别系统"', '深蓝背景，白色字体'],
                   ['数据看板', '设备号、工位名称、软件版本、当前时间、入库/出库总量、当班产量、当前装载数量、托盘号、产线状态', '实时显示系统运行数据'],
                   ['控制面板', '"手动运行"按钮和"手动停止"按钮', '手动运行: 黄灯亮 + 启动RFID盘存；手动停止: 停止RFID盘存 + 绿灯亮 + 显示标签'],
                   ['取标内容', '标签详细信息', '出入库完成时显示标签列表及写入校验结果(✓/✗)；手动停止时显示当前标签'],
                   ['系统日志', '通信日志', '实时滚动显示系统通信日志'],
               ])

    # ===== 9. 部署说明 =====
    add_heading_styled(doc, '9. 部署说明', 1)

    add_heading_styled(doc, '9.1 环境要求', 2)
    make_table(doc,
               ['项目', '要求'],
               [
                   ['操作系统', 'Linux (ARM/x86) / Windows / macOS'],
                   ['Python', '3.7+'],
                   ['依赖包', 'pyserial, paho-mqtt, openpyxl, tkinter (标准库)'],
               ])

    add_heading_styled(doc, '9.2 配置文件 (main.py 头部)', 2)
    cfg = '''# 串口配置
SERIAL_COM_IO = "/dev/ttyS0"              # IO模块串口
SERIAL_COM_RFID_READER = "/dev/ttysWK3"   # RFID读写器串口
SERIAL_COM_BARCODE_SCANNER = "/dev/ttyS1" # 条码扫描器串口

# MQTT配置 (MqttClient初始化参数)
broker = '192.168.3.83'
port = 1883

# TCP Server配置
host = '0.0.0.0'
port = 3000'''
    add_code_block(doc, cfg)

    add_heading_styled(doc, '9.3 启动方式', 2)
    doc.add_paragraph('1. 安装依赖: pip install -r requirements.txt')
    doc.add_paragraph('2. 根据实际硬件修改 main.py 头部的串口路径和网络配置')
    doc.add_paragraph('3. 运行: python main.py')
    doc.add_paragraph('4. 系统自动连接串口设备、MQTT Broker 和启动 TCP Server')

    add_heading_styled(doc, '9.4 启动后自动连接流程', 2)
    auto_text = """
启动顺序 (auto_connect):
  t+0s:  GUI界面加载
  t+2s:  MQTT客户端连接 (后台线程)
  t+3s:  IO模块串口连接 (后台线程)
  t+4s:  RFID读写器串口连接，固件启动，接收循环启动
  t+6s:  条码扫描器串口连接，接收循环启动

同时: TCP Socket Server 在 __init__ 中立即启动
"""
    add_code_block(doc, auto_text)

    # ===== 页脚 =====
    doc.add_paragraph()
    doc.add_paragraph('— 文档结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存
    output_path = '/Users/yli/Desktop/RFID产品/code/python/RFID_Detector_py/RFID_Detector_py/readme.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    build_document()
