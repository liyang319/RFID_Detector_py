"""生成 RFID项目需求规格书.docx，参考 软件需求说明书 OF V1.0.2.docx 格式"""
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_paragraph_run(para, text, font_name='宋体', font_size=Pt(10.5), bold=False, color=None, alignment=None):
    """向段落添加格式化文本"""
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        para.alignment = alignment
    return run


def add_normal_para(doc, text, font_name='宋体', font_size=Pt(10.5), indent=False, bold=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    add_paragraph_run(p, text, font_name=font_name, font_size=font_size, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    return p


def add_list_item(doc, text, font_name='宋体', font_size=Pt(10.5)):
    """添加列表项"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(21)
    add_paragraph_run(p, text, font_name=font_name, font_size=font_size)
    return p


def add_heading_1(doc, text):
    """添加一级标题"""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.font.bold = True
    return h


def add_heading_2(doc, text):
    """添加二级标题"""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.font.bold = True
    return h


def add_code_block(doc, text):
    """添加代码块（Topic/JSON）"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(42)
        add_paragraph_run(p, line, font_name='Consolas', font_size=Pt(8.5))


def add_code_block_no_indent(doc, text):
    """添加代码块（无额外缩进）"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_paragraph_run(p, line, font_name='Consolas', font_size=Pt(8.5))


def make_table(doc, headers, rows, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        add_paragraph_run(cell.paragraphs[0], h, font_name='宋体', font_size=Pt(9), bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            add_paragraph_run(cell.paragraphs[0], str(val), font_name='宋体', font_size=Pt(9))

    doc.add_paragraph()  # 表后空行
    return table


def add_page_number(doc):
    """添加页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 添加页码域
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)


def add_toc(doc, items):
    """添加目录"""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_paragraph_run(p, item, font_name='宋体', font_size=Pt(10.5))


def build_srs():
    doc = Document()

    # ===== 页面设置 =====
    for section in doc.sections:
        section.page_width = Cm(21)   # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 修改Normal样式默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面/标题页 =====
    # 空行
    for _ in range(6):
        doc.add_paragraph()

    # 软件名称
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_run(title_p, 'RFID标签识别系统', font_name='黑体', font_size=Pt(18), bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # 需求规格说明
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_run(sub_p, '软件需求规格说明书', font_name='黑体', font_size=Pt(16), bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # 版本信息
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_run(ver_p, '版本：RFID V1.0.0', font_name='宋体', font_size=Pt(12))

    doc.add_paragraph()
    ts_p = doc.add_paragraph()
    ts_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_run(ts_p, f'日期：{datetime.date.today().isoformat()}', font_name='宋体', font_size=Pt(12))

    doc.add_page_break()

    # ===== 版本说明 =====
    add_normal_para(doc, '版本说明：', bold=True)
    add_normal_para(doc, '当前版本：RFID V1.0.0，初始版本。', indent=True)

    doc.add_paragraph()

    # ===== 版本历史表 =====
    make_table(doc,
               ['序号', '变更章节', '变更说明', '变更日期'],
               [
                   ['1', '全部', '初始版本，基于RFID_Detector_py代码功能编写', datetime.date.today().isoformat()],
               ])

    doc.add_page_break()

    # ===== 目录 =====
    title_toc = doc.add_paragraph()
    title_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_run(title_toc, '目  录', font_name='黑体', font_size=Pt(14), bold=True)

    doc.add_paragraph()

    toc_items = [
        '1. 系统概述',
        '   1.1 软件标识',
        '   1.2 软件功能',
        '   1.3 硬件支撑',
        '   1.4 软件组成',
        '2. 用户界面',
        '   2.1 数据看板',
        '   2.2 取标内容',
        '   2.3 系统日志',
        '   2.4 控制面板',
        '3. 业务功能',
        '   3.1 入库识别',
        '   3.2 出库识别',
        '   3.3 手动运行/停止',
        '   3.4 紧急制动',
        '4. 标签写入与校验',
        '   4.1 写入数据来源',
        '   4.2 写标签流程',
        '   4.3 读标签验证',
        '   4.4 校验与上报',
        '5. 数据通信',
        '   5.1 TCP Socket Server',
        '   5.2 MQTT数据上报',
        '   5.3 IO模块串口通信',
        '   5.4 RFID读写器串口通信',
        '   5.5 条码扫描器串口通信',
        '6. 数据管理',
        '   6.1 标签数据查看',
        '   6.2 数据导出',
        '   6.3 日志记录',
        '7. 系统设置',
        '   7.1 串口配置',
        '   7.2 网络配置',
        '   7.3 写入数据配置',
        '8. 状态机与出入库逻辑',
        '   8.1 光栅状态定义',
        '   8.2 状态机流转',
        '   8.3 指示灯控制',
    ]
    add_toc(doc, toc_items)

    doc.add_page_break()

    # ===== 第1章 系统概述 =====
    add_heading_1(doc, '1. 系统概述')
    add_normal_para(doc, '软件名称：RFID标签识别系统')
    add_normal_para(doc, '软件功能：实现仓库出入库环节中RFID标签的自动写入、读取、校验与数据上报，配合光栅实现货物方向的自动识别，通过条码扫描进行货物数量比对，并通过MQTT和TCP与上位机系统进行数据交互。')
    add_normal_para(doc, '硬件支撑：RFID读写器(SFM2200)、IO模块(光栅×2、指示灯×3、蜂鸣器)、条码扫描器、工业计算机、电源等。')
    add_normal_para(doc, '软件组成：用户界面、出入库业务逻辑、标签写入与校验、数据通信(TCP/MQTT/串口)、数据管理、系统配置。')

    # ===== 第2章 用户界面 =====
    add_heading_1(doc, '2. 用户界面')
    add_normal_para(doc, '系统采用Tkinter图形框架，以工业风格深蓝色为主色调，界面支持垂直滚动，适配不同分辨率显示设备。界面从上至下分为以下几个区域：')

    add_heading_2(doc, '2.1 数据看板')
    add_normal_para(doc, '数据看板区域实时展示系统运行的关键指标，包含以下信息：')
    items = [
        '设备号：系统唯一标识符，默认"RFID-DETECTOR-001"；',
        '工位名称：可编辑文本输入框，默认"通道机-001"，用于标识当前工位；',
        '软件版本：当前软件版本号；',
        '当前时间：实时更新的系统时间（年-月-日 时:分:秒）；',
        '入库总量：累计完成的入库批次数；',
        '出库总量：累计完成的出库批次数；',
        '当班产量：当前班次的生产数量；',
        '当前装载数量：本次出入库识别到的RFID标签数量；',
        '托盘号：可编辑文本输入框，用于记录当前托盘编号；',
        '产线状态：正常(绿色)/异常(红色)指示灯及运行时间。',
    ]
    for item in items:
        add_list_item(doc, item)

    add_heading_2(doc, '2.2 取标内容')
    add_normal_para(doc, '取标内容区域为多行文本显示框，用于展示出入库完成后或手动停止后的标签详情。显示内容包括：')
    items2 = [
        '出入库方向（入库/出库）；',
        '写入校验结果：显示校验通过/失败及匹配数量；',
        '数量比对结果：显示标签数量与条码数量是否一致；',
        '每个标签的详细信息：校验标记(✓/✗)、产品名称、EPC、USER_DATA；',
        '条码列表：本次识别到的所有条码。',
    ]
    for item in items2:
        add_list_item(doc, item)

    add_heading_2(doc, '2.3 系统日志')
    add_normal_para(doc, '系统日志区域位于界面最下方，自动拉伸填充剩余空间。以"通信日志"为标签，实时滚动显示系统运行过程中的通信日志信息。日志格式为：[时间戳] 消息内容。日志自动限制行数，超出时自动清理旧内容。')

    add_heading_2(doc, '2.4 控制面板')
    add_normal_para(doc, '控制面板提供两个操作按钮：')
    items3 = [
        '"手动运行"按钮（绿色）：点击后切换运行状态。启动时：黄灯亮、绿灯灭、启动RFID读写器盘存循环(startloop)；停止时：停止RFID读写器盘存循环(stoploop)、绿灯亮、黄灯灭、将当前标签显示在取标内容窗口。',
        '"手动停止"按钮（红色，紧急制动）：点击后停止RFID读写器盘存循环(stoploop)、绿灯亮、黄灯灭、将当前标签显示在取标内容窗口，弹窗提示"数据已经上报！"。注意：手动停止不执行MQTT上报，仅显示标签。',
    ]
    for item in items3:
        add_list_item(doc, item)

    # ===== 第3章 业务功能 =====
    add_heading_1(doc, '3. 业务功能')

    add_heading_2(doc, '3.1 入库识别')
    add_normal_para(doc, '入库识别是指货物从入口方向（光栅1侧）进入通道，系统自动识别并写入标签、读取验证、上报数据的过程。')
    add_normal_para(doc, '入库流程如下：', bold=True)
    steps_in = [
        '1. 上位机通过TCP下发本次入库的写入数据（JSON格式：{"type":"rfid","data":[...]}）；',
        '2. 货物从入口进入，光栅1被遮挡 → 系统识别为入库开始（STATE_INBOUND_START）；',
        '3. 黄灯亮，绿灯灭；清空标签历史和条码缓存；',
        '4. 执行标签写入操作（优先使用TCP下发的pending_write_data，否则使用默认FIXED_USER_DATA），写入成功后启动RFID盘存以读回验证；',
        '5. 货物继续前进，出现以下两种路径之一：',
        '   Path1（标准路径）：光栅1和光栅2同时遮挡 → STATE_INBOUND_MIDDLE → 光栅2遮挡 → STATE_INBOUND_END；',
        '   Path2（快速通过）：光栅1松开后直接无遮挡 → STATE_INBOUND_END（跳过MIDDLE状态）；',
        '6. 到达END状态，光栅全部松开：绿灯亮，黄灯灭；停止RFID盘存；',
        '7. 获取本次所有条码，进行校验：比较读回的user_data与实际写入的actual_write_data是否一致；',
        '8. 通过MQTT上报标签数据、条码和校验结果；',
        '9. 通过TCP向所有已连接客户端发送通行完成消息：{"type":"pass","number":N}（N为标签数量）；',
        '10. 清空tag_history，状态机回到IDLE，等待下一次触发。',
    ]
    for step in steps_in:
        add_list_item(doc, step)

    add_heading_2(doc, '3.2 出库识别')
    add_normal_para(doc, '出库识别是指货物从出口方向（光栅2侧）进入通道，系统自动识别并写入标签、读取验证、上报数据的过程。')
    add_normal_para(doc, '出库流程如下：', bold=True)
    steps_out = [
        '1. 上位机通过TCP下发本次出库的写入数据；',
        '2. 货物从出口进入，光栅2被遮挡 → 系统识别为出库开始（STATE_OUTBOUND_START）；',
        '3. 黄灯亮，绿灯灭；清空标签历史和条码缓存；',
        '4. 执行标签写入操作（同入库）；',
        '5. 货物继续前进：',
        '   Path1（标准路径）：光栅1和光栅2同时遮挡 → STATE_OUTBOUND_MIDDLE → 光栅1遮挡 → STATE_OUTBOUND_END；',
        '   Path2（快速通过）：光栅2松开后直接无遮挡 → STATE_OUTBOUND_END（跳过MIDDLE状态）；',
        '6. 到达END状态后，流程同入库（校验、MQTT上报、TCP回传pass消息、回到IDLE）。',
    ]
    for step in steps_out:
        add_list_item(doc, step)

    add_heading_2(doc, '3.3 手动运行/停止')
    add_normal_para(doc, '"手动运行"按钮（toggle_production方法）用于手动控制系统的启停：')
    items_manual = [
        '启动（is_running=True）：黄灯亮，绿灯灭，调用rfid_reader_serial.startloop()启动RFID盘存；',
        '停止（is_running=False）：调用rfid_reader_serial.stoploop()停止RFID盘存，绿灯亮，黄灯灭，将当前标签显示在取标内容窗口。',
    ]
    for item in items_manual:
        add_list_item(doc, item)

    add_heading_2(doc, '3.4 紧急制动')
    add_normal_para(doc, '"手动停止"按钮（emergency_stop方法）用于紧急情况下的快速停止：')
    items_em = [
        '设置is_running=False；',
        '调用rfid_reader_serial.stoploop()停止RFID盘存；',
        '绿灯亮，黄灯灭；',
        '将当前标签显示在取标内容窗口；',
        '弹窗提示"数据已经上报！"。',
    ]
    for item in items_em:
        add_list_item(doc, item)

    # ===== 第4章 标签写入与校验 =====
    add_heading_1(doc, '4. 标签写入与校验')

    add_heading_2(doc, '4.1 写入数据来源')
    add_normal_para(doc, '标签写入数据支持两种来源，按优先级：')
    items_src = [
        '优先：上位机通过TCP Socket Server（端口3000）下发的JSON数据，格式为 {"type":"rfid","data":[整数数组]}。系统将data数组转为bytes，不足20字节时尾部补0x00，超过20字节时截取前20字节。此数据存储在pending_write_data中，不会因上报完成而被清除，仅在收到新TCP数据时覆盖。',
        '默认：当pending_write_data为None时，使用系统预设的FIXED_USER_DATA（20字节固定内容：11 22 33 44 55 66 77 88 99 00 AA AA BB BB CC CC DD DD EE EE）。',
    ]
    for item in items_src:
        add_list_item(doc, item)

    add_heading_2(doc, '4.2 写标签流程')
    add_normal_para(doc, '写标签由_execute_fixed_write()方法实现，在出入库START状态或MIDDLE状态（补充触发）时调用：')
    items_write = [
        '1. 设置write_in_progress=True，防止重复触发；',
        '2. 确定写入数据：write_data = pending_write_data（优先）或 FIXED_USER_DATA（默认）；',
        '3. 记录actual_write_data = write_data，用于后续校验比对；',
        '4. 调用RFIDReader_SFM2200.write_tag_with_userdata(write_data)写入标签；',
        '5. 写入成功：设置write_done=True，调用startloop()启动RFID盘存读取验证；',
        '6. 写入失败：自动重试一次；',
        '7. 两次均失败：设置write_done=False，仍启动startloop()读取标签原始数据。',
    ]
    for item in items_write:
        add_list_item(doc, item)

    add_heading_2(doc, '4.3 读标签验证')
    add_normal_para(doc, 'RFID读写器(SFM2200)持续盘存，通过_receive_loop接收标签上报数据。标签数据以FF 2B AA 00 00 00 96为包头，包含RSSI、天线号、User Data(20字节)、EPC等信息。')
    add_normal_para(doc, '标签解析后通过_add_serial_tag_to_history添加到tag_history中，基于EPC去重：')
    items_read = [
        '新EPC：直接添加到tag_history；',
        'EPC重复且write_done=True：覆盖已有标签的user_data（因写入操作只修改USER_DATA不修改EPC），确保校验使用的是写入后读回的最新数据；',
        'EPC重复且write_done=False：忽略，记录"检测到重复标签"。',
    ]
    for item in items_read:
        add_list_item(doc, item)

    add_heading_2(doc, '4.4 校验与上报')
    add_normal_para(doc, '在出入库完成（END状态或MIDDLE快速通过）或手动停止时，执行以下校验：')
    items_check = [
        '1. 对tag_history中每个标签，比较其读回的user_data（去除空格，转大写）与actual_write_data（十六进制大写）是否一致；',
        '2. 一致则标记write_verified=True（显示✓），不一致则标记write_verified=False（显示✗）；',
        '3. 比较标签数量与条码数量：一致显示"通过"，不一致显示"不一致"并提示具体数量；',
        '4. 构建包含标签列表、条码列表、校验结果、写入状态的report_data；',
        '5. 通过MQTT发布到rfid/data/{device_id}主题；',
        '6. 通过TCP向所有客户端发送{"type":"pass","number":N}；',
        '7. 清理tag_history，重置write_done=False，保留pending_write_data（不清除）。',
    ]
    for item in items_check:
        add_list_item(doc, item)

    # ===== 第5章 数据通信 =====
    add_heading_1(doc, '5. 数据通信')

    add_heading_2(doc, '5.1 TCP Socket Server')
    add_normal_para(doc, 'TCP Socket Server运行在端口3000（host=0.0.0.0），支持多个客户端并发连接。每个客户端由独立线程处理。')
    add_normal_para(doc, '接收消息格式（上位机→系统）：', bold=True)
    tcp_recv = '''{
    "type": "rfid",
    "data": [48, 50, 87, 68, 54, 0, 11, 0, 24, 88, 0, 12, 26, 1, 1, 0, 100, 0, 10]
}'''
    add_code_block(doc, tcp_recv)
    add_normal_para(doc, '说明：type固定为"rfid"，data为0-255的整数数组。系统自动将数组转为bytes并补齐/截断到20字节。')
    add_normal_para(doc, '此外还支持以下文本控制命令：')
    items_tcp = [
        '"stop"：触发紧急制动（emergency_stop）；',
        '"start"：触发运行状态切换（toggle_production）；',
        '"status"：返回当前运行状态和识别数量。',
    ]
    for item in items_tcp:
        add_list_item(doc, item)

    add_normal_para(doc, '发送消息格式（系统→上位机）：', bold=True)
    tcp_send = '''{
    "type": "pass",
    "number": 5
}'''
    add_code_block(doc, tcp_send)
    add_normal_para(doc, '说明：type固定为"pass"，number为本次识别到的RFID标签数量。在每次出入库完成时自动广播给所有已连接客户端。')

    add_heading_2(doc, '5.2 MQTT数据上报')
    add_normal_para(doc, 'MQTT客户端连接至Broker（默认192.168.3.83:1883），Client ID格式为"RFID_DETECTOR_{device_id}"。')
    add_normal_para(doc, '上报Topic：rfid/data/{device_id}')
    add_normal_para(doc, '上报数据格式：', bold=True)
    mqtt_fmt = '''{
    "cmd": "report_tags",
    "type": "inbound",
    "data": {
        "tags": [
            {
                "epc": "C090C000000A4B28",
                "tid": "",
                "user_data": "30 32 57 44 36 00 0B 00 ...",
                "rssi": -42.5,
                "antenna_num": 1,
                "pc": "04 FF",
                "product_name": "乳化炸药",
                "manufacturer": "默认生产企业",
                "license_number": "SC20240001",
                "production_date": "2024-08-15",
                "batch_number": "BATCH-...",
                "package_spec": "标准规格",
                "package_method": "箱装",
                "quantity": 1,
                "longitude": 116.3918,
                "latitude": 39.9798,
                "timestamp": "2026-05-29 12:00:00",
                "write_verified": true
            }
        ],
        "barcodes": ["6923456789012"],
        "validation": {
            "write_verified_count": 1,
            "write_total_count": 1,
            "errors": []
        },
        "write_success": true
    }
}'''
    add_code_block_no_indent(doc, mqtt_fmt)

    add_heading_2(doc, '5.3 IO模块串口通信')
    add_normal_para(doc, 'IO模块通过串口（默认/dev/ttyS0, 9600bps）通信，采用类Modbus RTU协议：')
    items_io = [
        '帧格式：FE + 命令(1字节) + 数据 + CRC16(MODBUS)',
        '读寄存器：read_register(cmd) — 发送 FE cmd 00 00 00 08 + CRC16，返回光栅状态；',
        '写寄存器：write_register(port, b_on) — port为寄存器地址(0x00红/0x02黄/0x04绿/0x06蜂鸣器)，b_on=True写0xFF(亮/响)，b_on=False写0x00(灭/停)，发送 FE 05 00 port data 00 + CRC16；',
        '光栅状态从地址0x02读取，bit0=光栅1(入口)，bit1=光栅2(出口)。组合值为：0x00无遮挡/0x01光栅1遮挡/0x02光栅2遮挡/0x03光栅1+2同时遮挡。',
    ]
    for item in items_io:
        add_list_item(doc, item)

    add_heading_2(doc, '5.4 RFID读写器串口通信')
    add_normal_para(doc, 'RFID读写器(SFM2200)通过串口（默认/dev/ttysWK3, 115200bps）通信，使用SFM2200私有协议。')
    items_rfid = [
        '指令帧格式：包含CRC16-CCITT校验，校验多项式为0x1021；',
        '标签盘存指令(startloop)：发送FF 1F AA...指令，读写器开始持续上报标签数据；',
        '停止盘存指令(stoploop)：发送FF 0E AA...指令，停止标签上报；',
        '写标签指令(write_tag_with_userdata)：将20字节userdata打包为31字节命令帧(11字节固定头+20字节userdata)，发送后等待响应；',
        '响应格式：FF 00 24 + 状态码(2字节，0x0000=成功)；',
        '标签上报数据：以FF 2B AA 00 00 00 96为包头，包含RSSI、天线号、User Data(20字节)、EPC等信息，详见第4.3节。',
    ]
    for item in items_rfid:
        add_list_item(doc, item)

    add_heading_2(doc, '5.5 条码扫描器串口通信')
    add_normal_para(doc, '条码扫描器通过串口（默认/dev/ttyS1, 9600bps）通信，接收ASCII格式的条码字符串。')
    items_bc = [
        '后台接收线程持续监听串口数据；',
        '收到数据后解码为ASCII字符串，过滤控制字符，验证长度(3-50位)；',
        '有效条码添加到条码队列（自动去重），并记录到CSV日志文件(barcode_log_YYYYMMDD.csv)；',
        '支持回调函数通知主程序有新条码到达；',
        'get_all_barcodes()方法获取并清空当前累积的所有条码。',
    ]
    for item in items_bc:
        add_list_item(doc, item)

    # ===== 第6章 数据管理 =====
    add_heading_1(doc, '6. 数据管理')

    add_heading_2(doc, '6.1 标签数据查看')
    add_normal_para(doc, '标签数据可通过以下方式查看：')
    items_view = [
        '实时查看：出入库完成后，"取标内容"区域显示本批次所有标签的EPC、USER_DATA、校验结果等；',
        '手动查看：点击"手动停止"后，取标内容显示当前累积的所有标签信息；',
        '历史记录：tag_history保存当前批次的标签列表，最大容量10000条。',
    ]
    for item in items_view:
        add_list_item(doc, item)

    add_heading_2(doc, '6.2 数据导出')
    add_normal_para(doc, '支持将tag_history中的标签数据导出为Excel(.xlsx)文件。导出字段包括：时间戳、EPC、TID、USER_DATA、RSSI、天线号、产品名称、生产企业、许可证号、生产日期、批号、包装规格、包装方式、数量、经度、纬度、解析状态。')

    add_heading_2(doc, '6.3 日志记录')
    add_normal_para(doc, '系统提供多种日志记录：')
    items_log = [
        '通信日志：UI系统日志窗口实时显示所有通信日志，带时间戳；',
        'RFID读写器日志：RFIDReader_SFM2200.log_data()记录RFID原始数据到rfid_reader_log.txt；',
        '条码日志：BarCodeScanner.log_barcode()记录条码到barcode_log_YYYYMMDD.csv。',
    ]
    for item in items_log:
        add_list_item(doc, item)

    # ===== 第7章 系统设置 =====
    add_heading_1(doc, '7. 系统设置')

    add_heading_2(doc, '7.1 串口配置')
    add_normal_para(doc, '系统串口配置在main.py文件头部定义：')
    serial_cfg = '''SERIAL_COM_IO = "/dev/ttyS0"              # IO模块（光栅+灯+蜂鸣器）串口
SERIAL_COM_RFID_READER = "/dev/ttysWK3"   # RFID读写器(SFM2200)串口
SERIAL_COM_BARCODE_SCANNER = "/dev/ttyS1" # 条码扫描器串口'''
    add_code_block(doc, serial_cfg)
    add_normal_para(doc, '可根据实际硬件连接修改上述路径。IO模块波特率为9600，RFID读写器波特率为115200，条码扫描器波特率为9600。')

    add_heading_2(doc, '7.2 网络配置')
    add_normal_para(doc, '网络配置包括MQTT和TCP两部分：')
    items_net = [
        'MQTT Broker：默认192.168.3.83:1883，在MqttClient初始化参数中配置；',
        'MQTT认证：username和password可在初始化参数中设置；',
        'TCP Server：默认监听0.0.0.0:3000，在TcpSocketServer初始化参数中配置。',
    ]
    for item in items_net:
        add_list_item(doc, item)

    add_heading_2(doc, '7.3 写入数据配置')
    add_normal_para(doc, '标签写入数据的默认值(FIXED_USER_DATA)在RFIDProductionSystem.__init__中定义（20字节）。当上位机未通过TCP下发写入数据时，系统使用此默认值。')

    # ===== 第8章 状态机与出入库逻辑 =====
    add_heading_1(doc, '8. 状态机与出入库逻辑')

    add_heading_2(doc, '8.1 光栅状态定义')
    make_table(doc,
               ['状态值', '含义', '说明'],
               [
                   ['0x00', '无遮挡', '通道内无货物'],
                   ['0x01', '光栅1遮挡', '入口端光栅被遮挡，货物从入口方向进入'],
                   ['0x02', '光栅2遮挡', '出口端光栅被遮挡，货物从出口方向进入'],
                   ['0x03', '光栅1+2同时遮挡', '货物在通道中间，两端光栅均被遮挡'],
               ])

    add_heading_2(doc, '8.2 状态机流转')
    add_normal_para(doc, '状态机定义了6个状态，通过光栅状态变化驱动流转：')
    make_table(doc,
               ['状态名', '值', '触发条件', '执行动作'],
               [
                   ['STATE_IDLE', '0', '初始状态/完成状态', '等待光栅触发'],
                   ['STATE_INBOUND_START', '1', 'IDLE + 光栅1遮挡(0x01)', '黄灯亮、清空缓存、写标签、启动RFID盘存'],
                   ['STATE_INBOUND_MIDDLE', '2', 'INBOUND_START + 两光栅同时遮挡(0x03)', '如写标签未触发则补充写入'],
                   ['STATE_INBOUND_END', '3', 'INBOUND_MIDDLE + 光栅2遮挡(0x02) 或 INBOUND_START + 直接无遮挡(0x00)', '停止RFID盘存、绿灯亮、校验、MQTT上报、TCP回传pass'],
                   ['STATE_OUTBOUND_START', '4', 'IDLE + 光栅2遮挡(0x02)', '黄灯亮、清空缓存、写标签、启动RFID盘存'],
                   ['STATE_OUTBOUND_MIDDLE', '5', 'OUTBOUND_START + 两光栅同时遮挡(0x03)', '如写标签未触发则补充写入'],
                   ['STATE_OUTBOUND_END', '6', 'OUTBOUND_MIDDLE + 光栅1遮挡(0x01) 或 OUTBOUND_START + 直接无遮挡(0x00)', '停止RFID盘存、绿灯亮、校验、MQTT上报、TCP回传pass'],
               ])

    add_normal_para(doc, '状态机使用确认机制防止误触发：光栅状态必须连续N次（可配置次数，如3次）读到同一状态才会被确认为有效状态转换，避免信号抖动。')

    add_normal_para(doc, '入库路径：', bold=True)
    add_normal_para(doc, 'Path1(标准): IDLE→(0x01)→START→(0x03)→MIDDLE→(0x02)→END→(0x00)→IDLE')
    add_normal_para(doc, 'Path2(快速): IDLE→(0x01)→START→(0x00)→END→(0x00)→IDLE')

    add_normal_para(doc, '出库路径：', bold=True)
    add_normal_para(doc, 'Path1(标准): IDLE→(0x02)→START→(0x03)→MIDDLE→(0x01)→END→(0x00)→IDLE')
    add_normal_para(doc, 'Path2(快速): IDLE→(0x02)→START→(0x00)→END→(0x00)→IDLE')

    add_heading_2(doc, '8.3 指示灯控制')
    make_table(doc,
               ['指示灯', '寄存器地址', '亮灯条件', '灭灯条件'],
               [
                   ['红灯', '0x00', '（预留）异常状态', '正常状态'],
                   ['黄灯', '0x02', '出入库START状态（货物进入通道，处理中）', '出入库END状态或手动停止'],
                   ['绿灯', '0x04', '系统就绪(IDLE)或出入库完成', '出入库START状态（货物处理中）'],
                   ['蜂鸣器', '0x06', '（预留）异常告警', '正常状态'],
               ])

    add_normal_para(doc, '指示灯通过write_register(port, b_on)控制，b_on=True发送0xFF（亮），b_on=False发送0x00（灭）。')

    # ===== 添加页码 =====
    add_page_number(doc)

    # 保存
    output_path = '/Users/yli/Desktop/RFID产品/code/python/RFID_Detector_py/RFID_Detector_py/RFID项目需求规格书.docx'
    doc.save(output_path)
    print(f'需求规格书已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    build_srs()
