"""生成 民用爆炸品企业服务管理平台需求说明书 v1.2"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def add_run(para, text, font_name='宋体', font_size=Pt(10.5), bold=False):
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    return run


def add_normal(doc, text, indent=False, bold=False):
    p = doc.add_paragraph()
    add_run(p, text, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    return p


def add_list(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(21)
    add_run(p, text)
    return p


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.font.bold = True


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.font.bold = True


def add_code(doc, text, indent=True):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if indent:
            p.paragraph_format.left_indent = Pt(42)
        add_run(p, line, font_name='Consolas', font_size=Pt(8.5))


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        add_run(cell.paragraphs[0], h, font_size=Pt(9), bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            add_run(cell.paragraphs[0], str(val), font_size=Pt(9))
    doc.add_paragraph()
    return table


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld)
        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        run2._r.append(instr)
        run3 = p.add_run()
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld2)


def build():
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.18)
        s.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tp, '民用爆炸品企业服务管理平台', font_name='黑体', font_size=Pt(18), bold=True)
    doc.add_paragraph()
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp, '需求说明书 v1.2', font_name='黑体', font_size=Pt(16), bold=True)
    doc.add_paragraph(); doc.add_paragraph()
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(vp, '版本：V1.2.0', font_size=Pt(12))
    doc.add_paragraph()
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(dp, f'日期：{datetime.date.today().isoformat()}', font_size=Pt(12))
    doc.add_page_break()

    # ===== 版本说明 =====
    add_normal(doc, '版本说明：', bold=True)
    add_normal(doc, '当前版本：V1.2.0。基于V1.1.0重构：将生产、出入库、销售、爆破、靶试、运输等独立管理模块整合为产品生命周期状态，统一纳入数据管理章节，突出以RFID标签为核心的产品信息数字化管理定位。', indent=True)
    doc.add_paragraph()
    make_table(doc,
        ['序号', '变更章节', '变更说明', '变更日期'],
        [
            ['1', '全部', '初始版本', '2026-06-10'],
            ['2', '3,4,5,12', '增加设备管理、标签打印管理、运输数据接入、运输监控', '2026-06-18'],
            ['3', '6,7-12', '重构：将独立业务管理模块整合为产品生命周期状态；明确平台以RFID标签为核心的数字化管理定位', '2026-06-22'],
            ['4', '4,5,6', '新增4.4标签信息上报第三方平台、5.8工厂原有数据导入(Excel)、6.6数据大屏接口；所有接口报文标注"待定"', datetime.date.today().isoformat()],
        ])
    doc.add_page_break()

    # ===== 目录 =====
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tp, '目  录', font_name='黑体', font_size=Pt(14), bold=True)
    doc.add_paragraph()
    toc = [
        '1. 系统概述', '   1.1 软件定位', '   1.2 软件目标', '   1.3 硬件支撑', '   1.4 软件组成',
        '2. 用户权限管理', '   2.1 管理员权限', '   2.2 企业用户权限', '   2.3 监管用户权限',
        '3. 终端设备管理', '   3.1 设备类型', '   3.2 设备注册与认证', '   3.3 设备状态监控', '   3.4 设备OTA升级',
        '4. 标签打印内容管理', '   4.1 打印模板管理', '   4.2 打印信息下发', '   4.3 打印记录追溯', '   4.4 标签信息上报第三方平台',
        '5. 数据采集与接入', '   5.1 产线通道机数据上报', '   5.2 廊机数据上报', '   5.3 手持设备数据接入', '   5.4 销售端数据接入', '   5.5 爆破作业数据接入', '   5.6 靶试实验（科研）数据接入', '   5.7 运输数据接入', '   5.8 工厂原有数据导入', '   5.9 第三方平台数据对接',
        '6. 数据管理', '   6.1 产品全生命周期状态', '   6.2 多维度数据查询', '   6.3 数据统计', '   6.4 报表生成', '   6.5 数据接口', '   6.6 数据大屏接口',
        '7. 系统对接', '   7.1 公安部管理平台对接', '   7.2 行业管理平台对接',
        '8. 告警与通知', '9. 系统设置', '10. 帮助',
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, item)
    doc.add_page_break()

    # ===== 1. 系统概述 =====
    add_h1(doc, '1. 系统概述')
    add_h2(doc, '1.1 软件定位')
    add_normal(doc, '软件名称：民用爆炸品企业服务管理平台')
    add_normal(doc, '软件定位：以RFID标签为核心的产品信息数据管理平台。平台不实现独立的生产、出入库、销售、爆破、运输等业务流程，而是通过终端设备采集和数据接口接入，将产品在各环节的状态信息统一汇聚，实现以EPC号为唯一标识的产品全生命周期数字化管理。')

    add_h2(doc, '1.2 软件目标')
    add_normal(doc, '构建民用爆炸品从生产下线到最终使用（爆破或科研）的全生命周期数据追溯平台。通过管理产线通道机、廊机、手持设备等各类终端，接收设备上报的数据，下发标签打印信息和写入指令，提供多维度数据查询、统计和报表，并通过标准API接口对接第三方监管平台。核心目标是实现爆炸品"来源可查、去向可追、状态可知、数据可溯"。')

    add_h2(doc, '1.3 硬件支撑')
    add_normal(doc, '云服务器、数据库服务器、产线通道机（RFID读写器+条码扫描器+IO模块）、廊机（库房门廊出入库识别设备）、手持终端（数据采集与上报）、数据大屏（实时监控展示）、网络设备等。')

    add_h2(doc, '1.4 软件组成')
    add_normal(doc, '用户权限管理、终端设备管理、标签打印内容管理、数据采集与接入、数据管理（核心）、系统对接、告警与通知、系统设置、帮助。平台以数据管理为核心，其他模块围绕数据采集、数据下发和数据输出展开。')

    # ===== 2. 用户权限管理 =====
    add_h1(doc, '2. 用户权限管理')
    add_h2(doc, '2.1 管理员权限')
    add_list(doc, '全权限：包括用户管理、系统配置、数据删除、平台对接配置等所有功能；')
    add_list(doc, '可创建、编辑、删除企业用户账号和监管用户账号。')
    add_h2(doc, '2.2 企业用户权限')
    add_list(doc, '数据查看：按条件查询本企业相关产品的全生命周期数据，查看统计报表；')
    add_list(doc, '打印管理：管理本企业各分厂/产线的标签打印内容模板和下发；')
    add_list(doc, '设备管理：查看本企业终端设备的在线状态和运行信息。')
    add_h2(doc, '2.3 监管用户权限')
    add_list(doc, '数据查看：查询所有企业上报的数据，按多维度检索；')
    add_list(doc, '统计报表：查看全域统计数据，生成监管报表；')
    add_list(doc, '告警接收：接收异常告警通知；')
    add_list(doc, '系统对接：管理第三方平台（公安部等）的对接配置。')

    # ===== 3. 终端设备管理 =====
    add_h1(doc, '3. 终端设备管理')
    add_normal(doc, '平台统一管理所有接入的终端设备，包括设备注册、认证、状态监控和OTA固件升级。')

    add_h2(doc, '3.1 设备类型')
    make_table(doc,
        ['设备类型', '应用场景', '主要功能', '通信协议'],
        [
            ['产线通道机', '安装在生产线末端', '标签写入、读取校验、条码扫描、数据上报', 'MQTT + TCP'],
            ['廊机', '安装在库房门廊处', '出入库识别、标签读取、条码扫描、数据上报', 'MQTT + TCP'],
            ['手持设备', '销售/爆破/运输现场', '标签读取、数据录入、定位上报', 'MQTT + HTTP REST'],
            ['数据大屏', '监控中心/调度室', '实时数据展示、告警展示、统计看板', 'HTTP REST / WebSocket'],
        ])

    add_h2(doc, '3.2 设备注册与认证')
    add_list(doc, '设备首次接入平台需进行注册，提交设备唯一标识（SN号/设备ID）、设备类型、所属企业/分厂、安装位置等信息；')
    add_list(doc, '平台审核通过后分配设备认证凭证（Token/证书），设备携带凭证进行后续通信；')
    add_list(doc, '支持设备信息编辑、禁用/启用、删除（注销）操作；')
    add_list(doc, '设备认证方式：基于Token的身份认证，每次MQTT连接和HTTP请求携带Token进行校验。')

    add_h2(doc, '3.3 设备状态监控')
    add_list(doc, '在线状态：实时展示各设备的在线/离线状态，离线超过阈值时间自动告警；')
    add_list(doc, '心跳检测：设备定时（默认30秒）向平台发送心跳消息，平台记录最近心跳时间；')
    add_list(doc, '设备信息查看：查看设备的基本信息、固件版本、最近上报数据时间、运行时长等；')
    add_list(doc, '设备列表筛选：按设备类型、所属企业、分厂、在线状态等条件筛选设备列表；')
    add_list(doc, '设备异常记录：记录设备异常事件（掉线、数据上报异常、硬件故障等），支持追溯查询。')

    add_h2(doc, '3.4 设备OTA升级')
    add_list(doc, '固件版本管理：平台管理各设备类型的固件版本，支持上传新版本固件包；')
    add_list(doc, '升级策略配置：支持立即升级、定时升级、分批次灰度升级；')
    add_list(doc, '升级目标选择：可按设备类型、所属企业、分厂、单个设备等多粒度选择升级目标；')
    add_list(doc, '升级进度监控：实时查看各设备的升级状态（等待中/下载中/安装中/成功/失败）；')
    add_list(doc, '升级失败回滚：升级失败时自动回滚到上一版本，并通知管理员；')
    add_list(doc, 'OTA升级Topic（待定）：firmware/upgrade/{device_id}（平台下发），firmware/status/{device_id}（设备上报状态）。')

    # ===== 4. 标签打印内容管理 =====
    add_h1(doc, '4. 标签打印内容管理')
    add_normal(doc, '平台统一管理不同产线的标签打印内容，支持按分厂、产线定点下发打印信息。标签打印内容在生产环节写入标签USER数据区，后续在通道机/廊机上被读取和校验。')

    add_h2(doc, '4.1 打印模板管理')
    add_list(doc, '模板创建：创建标签打印内容模板，定义各字段的内容和格式；')
    add_list(doc, '模板字段包括：产品类型、规格型号、生产厂家名称、生产厂家许可证号、生产日期、批号、包装规格、包装方式、数量、经纬度等；')
    add_list(doc, '模板支持变量替换：如{date}自动填充当前日期、{batch}自动填充批号等；')
    add_list(doc, '模板编辑与删除：支持模板的修改、复制、删除操作。')

    add_h2(doc, '4.2 打印信息下发')
    add_list(doc, '定点下发：选择目标分厂→产线，将指定的打印内容模板下发到对应产线的标签打印机或通道机；')
    add_list(doc, '下发内容格式：JSON格式，包含模板的所有字段值，通过MQTT下发到产线设备；')
    add_list(doc, '下发Topic（待定）：printer/content/{factory_id}/{line_id}（平台→设备）；')
    add_list(doc, '下发确认：设备收到打印信息后回复确认消息，平台记录下发状态；')
    add_list(doc, '批量下发：支持同时向多个分厂、多条产线下发相同或不同的打印内容；')
    add_list(doc, '下发历史：记录每次下发的目标、内容、时间、操作人，支持追溯查询。')
    add_normal(doc, '下发数据格式示例（待定）：', bold=True)
    print_fmt = '''{
    "type": "print_content",
    "factory_id": "F001",
    "line_id": "L03",
    "content": {
        "product_name": "乳化炸药",
        "package_spec": "Φ32mm×200g",
        "manufacturer": "XX化工有限公司",
        "license_number": "SC20240001",
        "production_date": "2026-06-11",
        "batch_number": "BATCH-20260611-001",
        "package_method": "箱装",
        "quantity": 24,
        "longitude": 116.3918,
        "latitude": 39.9798
    }
}'''
    add_code(doc, print_fmt)

    add_h2(doc, '4.3 打印记录追溯')
    add_list(doc, '每次打印下发自动记录：下发时间、目标分厂/产线、打印内容摘要、操作人员；')
    add_list(doc, '支持按时间范围、分厂、产线、产品类型等维度查询打印历史记录；')
    add_list(doc, '打印记录与后续生产数据通过EPC号关联，实现"打印内容下发→标签写入→读取校验→数据上报"全链路追溯。')

    add_h2(doc, '4.4 标签信息上报第三方平台')
    add_normal(doc, '支持将当前打印生产的标签信息，自动或手动上报到第三方数据平台（如公安部管理平台）。')
    add_list(doc, '上报触发方式：支持标签打印完成后自动上报，或通过平台界面手动选择批次上报；')
    add_list(doc, '上报数据内容：包含标签打印模板的全部字段（产品类型、规格型号、生产厂家、许可证号、生产日期、批号、数量等）及对应的EPC号、TID号列表；')
    add_list(doc, '上报方式：通过平台配置的第三方接口（RESTful API）发送JSON格式数据，接口地址和认证信息在"系统设置→数据对接配置"中配置；')
    add_list(doc, '上报状态跟踪：记录每次上报的时间、目标平台、数据量、成功/失败状态，失败时支持重试和告警；')
    add_list(doc, '上报数据格式（待定）：')
    report_fmt = '''{
    "type": "label_report",
    "report_time": "2026-06-23T10:00:00",
    "factory_id": "F001",
    "line_id": "L03",
    "batch_number": "BATCH-20260623-001",
    "labels": [{
        "epc": "C090C000000A4B28",
        "tid": "",
        "product_name": "乳化炸药",
        "package_spec": "Φ32mm×200g",
        "manufacturer": "XX化工有限公司",
        "license_number": "SC20240001",
        "production_date": "2026-06-23",
        "quantity": 24
    }]
}'''
    add_code(doc, report_fmt)

    # ===== 5. 数据采集与接入 =====
    add_h1(doc, '5. 数据采集与接入')
    add_normal(doc, '平台数据来源分为六种状态类型：生产、入库、出库、销售、爆破作业、靶试实验（科研）。其中生产数据由产线通道机通过MQTT自动上报，入库和出库数据由廊机上报，销售和爆破作业数据通过手持设备或Web端录入上报，靶试实验数据通过专用接口上报，运输数据由手持设备或车载终端实时上报。')

    add_h2(doc, '5.1 产线通道机数据上报（生产）')
    add_normal(doc, '通道机（RFID标签识别系统）在完成标签写入、读取校验后，通过MQTT协议向平台上报数据。')
    add_normal(doc, 'MQTT上报Topic（待定）：rfid/data/{device_id}', bold=True)
    add_normal(doc, '上报数据格式（待定）：', bold=True)
    mqtt_fmt = '''{
    "cmd": "report_tags",
    "type": "inbound",
    "data": {
        "tags": [{
            "epc": "C090C000000A4B28",
            "tid": "",
            "user_data": "30 32 57 44 36 00 ...",
            "rssi": -42.5,
            "antenna_num": 1,
            "pc": "04 FF",
            "product_name": "乳化炸药",
            "manufacturer": "XX化工有限公司",
            "license_number": "SC20240001",
            "production_date": "2024-08-15",
            "batch_number": "BATCH-...",
            "package_spec": "标准规格",
            "package_method": "箱装",
            "quantity": 1,
            "longitude": 116.3918,
            "latitude": 39.9798,
            "timestamp": "2026-06-05 12:00:00",
            "write_verified": true
        }],
        "barcodes": ["6923456789012"],
        "validation": {
            "write_verified_count": 1,
            "write_total_count": 1,
            "errors": []
        },
        "write_success": true
    }
}'''
    add_code(doc, mqtt_fmt)
    add_normal(doc, '通道机也通过TCP发送实时消息（待定）（report_rfid / report_barcode / cargo_in / cargo_out / pass），平台可建立TCP连接接收。')
    add_normal(doc, '生产端数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明', '来源'],
        [
            ['tid', 'string', '标签TID号', 'RFID读写器读取'],
            ['epc', 'string', '标签EPC号（唯一标识）', 'RFID读写器读取'],
            ['user_data', 'string', '用户数据区内容（十六进制）', '上位机下发→通道机写入'],
            ['longitude', 'float', '经度', '系统配置/通道机位置'],
            ['latitude', 'float', '纬度', '系统配置/通道机位置'],
            ['timestamp', 'datetime', '读取时间戳', '系统自动生成'],
            ['product_name', 'string', '产品类型（如乳化炸药）', '上位机下发'],
            ['package_spec', 'string', '规格型号', '上位机下发'],
            ['manufacturer', 'string', '生产厂家名称', '上位机下发'],
            ['license_number', 'string', '生产厂家许可证号', '上位机下发'],
            ['production_date', 'string', '生产日期', '上位机下发'],
            ['batch_number', 'string', '批号', '系统自动/上位机下发'],
            ['package_method', 'string', '包装方式', '上位机下发'],
            ['quantity', 'int', '数量', '上位机下发'],
            ['write_verified', 'bool', '标签写入校验是否通过', '通道机自动判定'],
        ])

    add_h2(doc, '5.2 廊机数据上报（入库/出库）')
    add_normal(doc, '廊机安装在库房门廊处，自动识别经过门廊的货物标签（入库或出库），并通过MQTT向平台上报数据。')
    add_list(doc, '廊机识别方向：通过内置光栅或红外传感器判断货物移动方向（入库/出库），与产线通道机状态机逻辑一致；')
    add_list(doc, '上报数据格式（待定）：与产线通道机相同，通过MQTT report_tags上报，data_type为"inbound"或"outbound"；')
    add_list(doc, '廊机也通过TCP发送实时消息（待定）（cargo_in / cargo_out / report_rfid / report_barcode），平台可建立TCP连接接收。')

    add_h2(doc, '5.3 手持设备数据接入')
    add_normal(doc, '手持设备用于销售现场、爆破作业现场、运输途中等场景的数据采集和上报。')
    add_list(doc, '功能包括：RFID标签读取、条码扫描、GPS定位、数据录入、拍照取证；')
    add_list(doc, '通信方式：通过4G/5G/WiFi网络，使用MQTT协议上报数据，使用HTTP REST进行配置同步；')
    add_list(doc, '上报数据类型：销售确认、爆破作业登记、运输状态更新、异常事件报告；')
    add_list(doc, '离线模式：网络不可用时本地缓存数据，恢复网络后自动补传。')

    add_h2(doc, '5.4 销售端数据接入')
    add_normal(doc, '销售端数据由销售企业在平台上录入或通过手持设备上报。数据作为产品生命周期"已销售"状态的记录。')
    add_normal(doc, '销售端数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['epc', 'string', '产品EPC号'],
            ['longitude', 'float', '销售地点经度'],
            ['latitude', 'float', '销售地点纬度'],
            ['timestamp', 'datetime', '销售时间'],
            ['product_name', 'string', '产品类型'],
            ['package_spec', 'string', '规格型号'],
            ['sale_company', 'string', '销售公司名称'],
            ['buyer_name', 'string', '购买方名称'],
            ['buyer_license', 'string', '购买方许可证号'],
        ])

    add_h2(doc, '5.5 爆破作业数据接入')
    add_normal(doc, '爆破作业数据由爆破公司在作业现场通过手持设备上报。数据作为产品生命周期"已使用（爆破）"状态的记录。')
    add_normal(doc, '爆破作业数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['epc', 'string', '产品EPC号'],
            ['longitude', 'float', '爆破作业地点经度'],
            ['latitude', 'float', '爆破作业地点纬度'],
            ['timestamp', 'datetime', '爆破作业时间'],
            ['product_name', 'string', '产品类型'],
            ['package_spec', 'string', '规格型号'],
            ['blast_company', 'string', '爆破公司名称'],
            ['blast_license', 'string', '爆破作业许可证号'],
            ['quantity', 'int', '爆破使用数量'],
            ['blast_method', 'string', '爆破方式'],
            ['blast_result', 'string', '爆破结果'],
        ])

    add_h2(doc, '5.6 靶试实验（科研）数据接入')
    add_normal(doc, '科研/靶试实验数据由实验单位上报。数据作为产品生命周期"已使用（科研）"状态的记录。')
    add_normal(doc, '科研数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['epc', 'string', '产品EPC号'],
            ['longitude', 'float', '实验地点经度'],
            ['latitude', 'float', '实验地点纬度'],
            ['timestamp', 'datetime', '实验时间'],
            ['product_name', 'string', '产品类型'],
            ['package_spec', 'string', '规格型号'],
            ['research_unit', 'string', '科研单位名称'],
            ['test_type', 'string', '实验类型（靶试/性能测试等）'],
            ['test_result', 'string', '实验结果'],
        ])

    add_h2(doc, '5.7 运输数据接入')
    add_normal(doc, '运输过程中的位置和状态数据由手持设备或车载终端实时上报，记录产品的流转轨迹。')
    add_normal(doc, '运输数据字段：', bold=True)
    make_table(doc,
        ['字段名', '类型', '说明'],
        [
            ['transport_id', 'string', '运输任务编号'],
            ['epc_list', 'array', '运输产品EPC号列表'],
            ['vehicle_plate', 'string', '运输车辆车牌号'],
            ['driver_name', 'string', '驾驶员姓名'],
            ['longitude', 'float', '当前位置经度（实时上报）'],
            ['latitude', 'float', '当前位置纬度（实时上报）'],
            ['speed', 'float', '当前速度（km/h）'],
            ['timestamp', 'datetime', '位置上报时间'],
            ['departure_place', 'string', '出发地'],
            ['destination', 'string', '目的地'],
            ['transport_status', 'string', '运输状态：出发/在途/到达/异常'],
        ])
    add_list(doc, '运输数据上报Topic（待定）：transport/data/{transport_id}（MQTT）；')
    add_list(doc, '上报频率：正常行驶时每30秒上报一次位置，异常事件（停车超时、偏离路线等）立即上报。')

    add_h2(doc, '5.8 工厂原有数据导入')
    add_normal(doc, '支持导入工厂原有系统中的产品数据，以Excel文件（.xlsx/.xls）的形式批量导入平台，实现历史数据的统一管理。')
    add_list(doc, '导入文件格式：Excel (.xlsx/.xls)，支持单个文件或批量文件夹导入；')
    add_list(doc, '导入数据字段映射（待定）：平台提供字段映射配置界面，将Excel列名映射到平台标准数据字段（EPC号、TID号、产品类型、规格型号、生产厂家、许可证号、生产日期、批号等），映射关系可保存为模板复用；')
    add_list(doc, '数据校验：导入时自动校验数据格式完整性、EPC号唯一性、必填字段是否缺失，校验不通过的数据行生成错误报告，支持下载查看和修正后重新导入；')
    add_list(doc, '重复处理策略：支持"跳过重复"（EPC号已存在则跳过）、"覆盖更新"（EPC号已存在则更新字段）、"仅导入新数据"三种策略；')
    add_list(doc, '导入进度：显示导入进度条、已导入/总计/跳过/错误数量统计；')
    add_list(doc, '导入记录：保存每次导入的历史记录（文件名、导入时间、操作人、导入数量、成功/失败/跳过数量），支持追溯查询。')
    add_normal(doc, '导入数据字段映射表示例（待定）：', bold=True)
    import_fmt = '''{
    "import_config": {
        "source": "excel",
        "field_mapping": {
            "epc": "EPC编码",
            "tid": "TID号",
            "product_name": "产品名称",
            "package_spec": "规格型号",
            "manufacturer": "生产企业",
            "license_number": "许可证号",
            "production_date": "生产日期",
            "batch_number": "批号",
            "quantity": "数量"
        },
        "duplicate_policy": "skip"
    }
}'''
    add_code(doc, import_fmt)

    add_h2(doc, '5.9 第三方平台数据对接')
    add_list(doc, '公安部管理平台对接：按《民用爆炸物品信息管理条例》要求，将产品生产、流通、使用数据同步上报至公安部管理平台。')
    add_list(doc, '行业管理平台对接：支持与行业协会或其他监管部门的平台进行数据交换。')
    add_list(doc, '对接方式：通过标准API接口（RESTful/WebService）或消息队列（MQTT/Kafka）进行数据同步。')
    add_list(doc, '对接数据格式：采用JSON格式，按对接平台要求进行字段映射和转换。')

    # ===== 6. 数据管理（核心） =====
    add_h1(doc, '6. 数据管理（核心）')
    add_normal(doc, '数据管理是平台的核心功能。平台不实现独立的生产管理、出入库管理、销售管理、爆破作业管理、靶试实验管理、运输监控管理等业务流程模块，而是将上述所有环节作为产品生命周期中的状态节点，统一在数据管理中进行查询、统计和分析。平台通过终端设备上报和数据接口接入，自动汇聚产品在各环节的状态信息，形成以EPC号为唯一标识的完整数据链。')

    add_h2(doc, '6.1 产品全生命周期状态')
    add_normal(doc, '以EPC号为核心标识，追踪产品从生产下线到最终使用的完整状态流转。每个产品在其生命周期中经历以下状态：')
    add_normal(doc, '生命周期状态定义：', bold=True)
    make_table(doc,
        ['状态', '状态标识', '触发条件', '记录的关键数据', '数据来源'],
        [
            ['已生产', 'produced', '产线通道机完成标签写入并上报', 'tid, epc, 产品类型, 规格型号, 生产厂家, 许可证号, 生产日期, 批号, 经纬度, 写入校验结果', '产线通道机 MQTT上报'],
            ['已入库', 'stored', '廊机检测入库方向并上报', 'epc, 入库时间, 仓库名称, 经纬度', '廊机 MQTT上报'],
            ['已出库', 'shipped', '廊机检测出库方向并上报', 'epc, 出库时间, 仓库名称, 出库去向, 经纬度', '廊机 MQTT上报'],
            ['运输中', 'transporting', '手持设备/车载终端上报运输数据', 'epc, 运输任务号, 车辆信息, 实时位置, 速度, 轨迹', '手持设备 MQTT上报'],
            ['已销售', 'sold', '销售企业上报销售数据', 'epc, 销售时间, 销售公司, 购买方信息, 经纬度', 'Web端/手持设备录入'],
            ['已使用（爆破）', 'used_blast', '爆破公司上报爆破作业数据', 'epc, 爆破时间, 爆破公司, 爆破地点, 使用量, 爆破结果, 经纬度', '手持设备上报'],
            ['已使用（科研）', 'used_research', '科研单位上报靶试实验数据', 'epc, 实验时间, 科研单位, 实验类型, 实验结果, 经纬度', '专用接口上报'],
            ['异常', 'abnormal', '数据校验失败/告警触发', 'epc, 异常类型, 异常时间, 异常描述', '平台自动判定'],
        ])
    add_normal(doc, '状态流转规则：')
    add_list(doc, '产品状态按时间顺序自动流转，后续状态的时间戳必须晚于前一状态；')
    add_list(doc, '同一EPC的同一状态不可重复记录（如已入库的产品不可再次入库）；')
    add_list(doc, '状态不可逆（如已出库的产品不可回退到已入库状态）；')
    add_list(doc, '异常状态可由任何正常状态触发，不影响正常状态链。')

    add_h2(doc, '6.2 多维度数据查询')
    add_normal(doc, '以产品生命周期状态数据为基础，支持以下维度的灵活查询：')
    make_table(doc,
        ['查询维度', '说明', '示例'],
        [
            ['EPC号', '精确查询单个产品的完整生命周期', '输入EPC号，展示该产品所有状态记录和时间线'],
            ['TID号', '通过TID查询关联产品', '输入TID号查询'],
            ['时间范围', '按起止日期筛选', '查询2026-06-01至2026-06-30的所有记录'],
            ['状态类型', '按生命周期状态筛选', '查询所有"已入库"或"运输中"的产品'],
            ['产品类型', '按产品名称筛选', '查询所有"乳化炸药"产品的记录'],
            ['规格型号', '按规格型号筛选', '查询"Φ32mm×200g"规格的所有记录'],
            ['生产企业', '按生产厂家或许可证号筛选', '查询XX化工有限公司的所有产品'],
            ['批号', '按生产批号查询', '查询BATCH-20260611-001批次的所有产品'],
            ['地理位置', '按行政区域或经纬度范围筛选', '查询河北省范围内所有产品的当前状态'],
            ['仓库', '按仓库筛选', '查询某仓库当前库存产品列表'],
            ['运输任务', '按运输任务编号查询', '查询某次运输任务中所有产品的实时状态'],
            ['设备ID', '按上报设备筛选', '查询某台通道机/廊机上报的所有数据'],
            ['组合查询', '上述维度的任意组合（AND/OR）', '查询XX企业2026年6月生产的已出库乳化炸药'],
        ])
    add_normal(doc, '查询结果支持列表展示和详情查看。列表展示关键字段摘要，点击单条记录可查看该产品完整的生命周期时间线和每个状态节点的详细数据。')

    add_h2(doc, '6.3 数据统计')
    add_normal(doc, '基于产品生命周期状态数据，提供以下统计维度：')
    add_list(doc, '按状态统计：统计各生命周期状态（已生产/已入库/已出库/运输中/已销售/已使用/异常）的产品数量和占比；')
    add_list(doc, '按时间统计：按日/周/月/季/年统计各产品类型的产量、出入库量、销售量、使用量趋势；')
    add_list(doc, '按企业统计：按生产企业、销售公司、爆破公司分别统计其产品的数量和状态分布；')
    add_list(doc, '按地域统计：以地图热力图展示产品在各地的生产、存储、使用分布情况；')
    add_list(doc, '按产品类型统计：统计各产品类型、规格型号的数量分布和占比；')
    add_list(doc, '校验统计：统计标签写入校验通过率、失败原因分类分析；')
    add_list(doc, '库存快照：指定时间点，统计各仓库各类产品的实时库存量；')
    add_list(doc, '异常统计：统计异常状态触发次数、类型分布、处理时长。')

    add_h2(doc, '6.4 报表生成')
    add_list(doc, '支持按日/周/月/季/年自动或手动生成各类统计报表；')
    add_list(doc, '报表类型：产品状态汇总报表、产品流向追踪报表、库存快照报表、异常分析报表、企业生产/销售/使用统计报表；')
    add_list(doc, '报表格式：PDF、Excel (.xlsx)、CSV；')
    add_list(doc, '支持报表模板自定义，用户可选择报表包含的字段、统计维度和时间范围；')
    add_list(doc, '支持报表自动定时生成和邮件发送。')

    add_h2(doc, '6.5 数据接口')
    add_normal(doc, '平台提供标准RESTful API接口，供第三方系统查询和对接产品生命周期数据。接口设计遵循以下原则：')
    add_list(doc, '以EPC号为核心查询键，支持单条查询和批量查询；')
    add_list(doc, '支持按状态、时间范围、企业等维度筛选；')
    add_list(doc, '返回JSON格式数据，包含产品完整生命周期状态链；')
    add_list(doc, '支持分页查询和大量数据导出；')
    add_list(doc, '接口认证方式（待定）：API Key + Token。')
    add_normal(doc, '主要API端点（待定）：', bold=True)
    make_table(doc,
        ['端点', '方法', '说明'],
        [
            ['/api/v1/products/{epc}', 'GET', '查询单个产品的完整生命周期数据'],
            ['/api/v1/products', 'GET', '按条件查询产品列表（支持多维度筛选参数）'],
            ['/api/v1/products/batch', 'POST', '批量查询多个EPC号的产品数据'],
            ['/api/v1/statistics', 'GET', '按指定维度和时间范围获取统计数据'],
            ['/api/v1/products/{epc}/timeline', 'GET', '获取单个产品的状态变更时间线'],
            ['/api/v1/devices/{device_id}/data', 'GET', '查询指定设备上报的历史数据'],
        ])

    add_h2(doc, '6.6 数据大屏接口')
    add_normal(doc, '平台提供专用的数据大屏接口，支持数据大屏设备通过WebSocket或HTTP REST实时获取产品信息的多维度展示数据。')
    add_normal(doc, '数据大屏接口（待定）支持以下功能：', bold=True)
    add_list(doc, '产品生命周期状态实时概览：各状态（已生产/已入库/已出库/运输中/已销售/已使用/异常）的产品数量统计，支持数字和图表展示；')
    add_list(doc, '地域分布热力图：以地图形式展示产品在全国/全省各地的生产、存储、使用分布情况，支持按产品类型筛选；')
    add_list(doc, '实时数据滚动：最新上报的产品数据以滚动列表形式实时展示（EPC号、产品类型、状态、时间、位置）；')
    add_list(doc, '设备在线状态监控：各类型终端设备的在线/离线数量统计，离线设备列表；')
    add_list(doc, '告警实时推送：最新的告警信息实时弹窗展示（告警类型、时间、设备、详情）；')
    add_list(doc, '产量/出入库量趋势图：按小时/日/周展示产量和出入库量的趋势折线图或柱状图；')
    add_list(doc, '运输监控地图：在地图上实时展示所有在途运输车辆的位置、轨迹和状态；')
    add_list(doc, '数据大屏支持自动轮播切换多个展示页面（可配置轮播间隔），支持手动切换页面。')
    add_normal(doc, '数据大屏主要接口端点（待定）：', bold=True)
    make_table(doc,
        ['端点', '方法', '说明'],
        [
            ['/api/v1/dashboard/overview', 'GET', '获取实时概览数据（各状态数量、设备在线数等）'],
            ['/api/v1/dashboard/geo', 'GET', '获取地域分布数据（热力图坐标点）'],
            ['/api/v1/dashboard/trend', 'GET', '获取趋势统计数据（产量/出入库量趋势）'],
            ['/api/v1/dashboard/alerts', 'GET', '获取最新告警列表'],
            ['/api/v1/dashboard/transport', 'GET', '获取在途运输车辆实时位置和状态'],
            ['/ws/dashboard/live', 'WebSocket', '实时数据推送（新上报数据、告警推送）'],
        ])

    # ===== 7. 系统对接 =====
    add_h1(doc, '7. 系统对接')

    add_h2(doc, '7.1 公安部管理平台对接')
    add_normal(doc, '按照GA/T 1225-2015《民用爆炸物品信息管理系统数据交换标准》等相关行业标准，实现与公安部民爆物品管理平台的数据对接。')
    add_list(doc, '数据同步内容：产品全生命周期状态数据（生产、出入库、销售、爆破作业、运输）；')
    add_list(doc, '数据同步方式：通过公安部指定的数据交换接口（API/中间库）定时或实时同步；')
    add_list(doc, '数据同步频率：支持实时同步和定时批量同步（可配置间隔）；')
    add_list(doc, '同步状态监控：记录每次同步的时间、数据量、成功/失败状态，失败时自动重试并告警。')

    add_h2(doc, '7.2 行业管理平台对接')
    add_list(doc, '支持与省级/市级民爆行业监管平台的数据对接；')
    add_list(doc, '支持与企业ERP/WMS系统的数据集成；')
    add_list(doc, '提供标准RESTful API接口（见6.5节）供第三方系统调用查询。')

    # ===== 8. 告警与通知 =====
    add_h1(doc, '8. 告警与通知')
    add_normal(doc, '平台对以下异常情况进行告警：')
    make_table(doc,
        ['告警类型', '触发条件', '告警级别', '通知方式'],
        [
            ['标签写入失败', '通道机上报write_success=false', '高', '平台弹窗 + 短信/邮件通知'],
            ['标签校验失败', 'write_verified_count < write_total_count', '高', '平台弹窗 + 短信/邮件通知'],
            ['数量不一致', '标签数量 ≠ 条码数量', '中', '平台弹窗'],
            ['设备离线', '终端设备心跳超时', '高', '平台弹窗 + 短信通知管理员'],
            ['运输路线偏离', '车辆偏离预设路线超过阈值', '高', '平台弹窗 + 短信通知驾驶员和管理员'],
            ['运输超时停留', '非目的地停留超过阈值', '高', '平台弹窗 + 短信通知驾驶员'],
            ['数据异常', '关键字段缺失或格式错误', '低', '平台弹窗'],
            ['第三方同步失败', '向公安部/行业平台同步数据失败', '高', '平台弹窗 + 短信通知'],
        ])
    add_normal(doc, '告警记录管理：所有告警记录保存至数据库，支持按时间、类型、级别查询；告警处理流程：产生→确认→处理→关闭，记录处理人和处理时间。')

    # ===== 9. 系统设置 =====
    add_h1(doc, '9. 系统设置')
    add_list(doc, '用户管理：创建、编辑、删除用户账号，分配角色权限；')
    add_list(doc, '企业管理：管理生产企业、销售公司、爆破公司、科研单位的基本信息（名称、许可证号、地址等）；')
    add_list(doc, '仓库管理：管理仓库信息（名称、位置、容量等）；')
    add_list(doc, '产品类型管理：管理产品分类和规格型号字典；')
    add_list(doc, '告警规则配置：配置各类告警的触发条件、级别和通知方式；')
    add_list(doc, '数据对接配置：配置与公安部平台、行业平台的对接参数（接口地址、认证信息、同步频率等）；')
    add_list(doc, '界面设置：界面风格、字体、主题切换；')
    add_list(doc, '系统日志：记录用户操作日志和系统运行日志，支持日志查询和导出。')

    # ===== 10. 帮助 =====
    add_h1(doc, '10. 帮助')
    add_list(doc, '版本说明：当前版本V1.2.0，后续版本更新记录；')
    add_list(doc, '使用说明：平台各模块的操作手册，含图文教程；')
    add_list(doc, 'API文档：平台对外接口的技术文档（见6.5节），含请求/响应示例；')
    add_list(doc, '常见问题（FAQ）：常见问题及解答；')
    add_list(doc, '在线支持：联系方式、工单提交。')

    add_page_numbers(doc)

    path = '民用爆炸品企业服务管理平台需求说明书 v1.2.docx'
    doc.save(path)
    print(f'文档已生成: {path}')


if __name__ == '__main__':
    build()
